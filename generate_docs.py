"""
generate_docs.py
Generates two Word documents:
  1. tech_brief.docx            — graphically designed 1-pager
  2. validation_questions.docx  — validator questionnaire
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


# ── Low-level helpers ──────────────────────────────────────────────────────────

def _hex(h):
    b = bytes.fromhex(h)
    return RGBColor(b[0], b[1], b[2])


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def no_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _get_tblPr(tbl):
    tbl_el = tbl._tbl
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    return tblPr


def no_table_borders(tbl):
    tblPr = _get_tblPr(tbl)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_table_width(tbl, cm_width):
    tblPr = _get_tblPr(tbl)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(cm_width * 567)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def cell_margin(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def divider(doc, color="2c3e50"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# 1. TECH BRIEF  (graphically designed, 1 page)
# ══════════════════════════════════════════════════════════════════════════════

# Brand palette
C_NAVY    = "1f3864"
C_BLUE    = "2471a3"
C_STEEL   = "adb5bd"
C_WHITE   = "ffffff"
C_LIGHT   = "eaf0fb"
C_DARK    = "2c3e50"

# Model chip colours (matching app MODEL_COLORS)
MODEL_CHIPS = [
    ("GRU",               "1f77b4"),
    ("LSTM",              "ff7f0e"),
    ("SimpleRNN",         "2ca02c"),
    ("Linear Regression", "d62728"),
    ("ARIMA",             "9467bd"),
]


def _zero_para(doc):
    """Insert a zero-height paragraph to separate adjacent tables."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "2")
    szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "2")
    rPr = OxmlElement("w:rPr")
    rPr.append(sz); rPr.append(szCs)
    pPr.append(rPr)


def add_col_divider(cell, color="bdc3c7", sz=6):
    """Thin right border on a cell, used as column separator."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    right_el = OxmlElement("w:right")
    right_el.set(qn("w:val"), "single")
    right_el.set(qn("w:sz"), str(sz))
    right_el.set(qn("w:space"), "0")
    right_el.set(qn("w:color"), color)
    tcBorders.append(right_el)
    tcPr.append(tcBorders)


def build_tech_brief():
    doc = Document()

    # ── Page setup: narrow margins, Letter (21.59 × 27.94 cm) ─────────────────
    for sec in doc.sections:
        sec.page_width    = Cm(21.59)
        sec.page_height   = Cm(27.94)
        sec.top_margin    = Cm(1.1)
        sec.bottom_margin = Cm(1.1)
        sec.left_margin   = Cm(1.5)
        sec.right_margin  = Cm(1.5)

    usable_cm = 21.59 - 3.0   # 18.59 cm

    # ══════════════════════════════════════════════════════════════════════════
    # BLOCK 1 — HEADER BANNER
    # ══════════════════════════════════════════════════════════════════════════
    hdr = doc.add_table(rows=1, cols=1)
    no_table_borders(hdr)
    hdr_c = hdr.rows[0].cells[0]
    set_cell_bg(hdr_c, C_NAVY)
    no_cell_borders(hdr_c)
    cell_margin(hdr_c, top=90, bottom=90, left=220, right=220)

    p_eye = hdr_c.paragraphs[0]
    p_eye.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eye.paragraph_format.space_before = Pt(0)
    p_eye.paragraph_format.space_after  = Pt(3)
    r_eye = p_eye.add_run("■  TECHNICAL BRIEF  ■")
    r_eye.font.size = Pt(7); r_eye.bold = True
    r_eye.font.color.rgb = _hex(C_STEEL)

    p_title = hdr_c.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(3)
    r_title = p_title.add_run(
        "Evaluating the Performance of Gated Recurrent Units\n"
        "for Multi-Parameter Weather Forecasting"
    )
    r_title.bold = True; r_title.font.size = Pt(14)
    r_title.font.color.rgb = _hex(C_WHITE)

    p_auth = hdr_c.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_auth.paragraph_format.space_before = Pt(0)
    p_auth.paragraph_format.space_after  = Pt(0)
    r_auth = p_auth.add_run(
        "University of the Cordilleras  ·  College of CITCS  ·  B.S. Computer Science  ·  "
        "Castro, E. M.  ·  Millan, D. J. S.  ·  Mondoñedo, J. E."
    )
    r_auth.font.size = Pt(8); r_auth.font.color.rgb = _hex("adc8e0")

    # ══════════════════════════════════════════════════════════════════════════
    # BLOCK 2 — HIGHLIGHT STATS ROW (4 colored boxes)
    # ══════════════════════════════════════════════════════════════════════════
    _zero_para(doc)
    stats = [
        ("5",  "MODELS\nEVALUATED",       "2471a3"),
        ("4",  "WEATHER\nVARIABLES",      "1a7a4a"),
        ("2",  "WEATHER\nSTATIONS",       "7d3c98"),
        ("5+", "YEARS OF\nTRAINING DATA", "c0392b"),
    ]
    stat_tbl = doc.add_table(rows=1, cols=len(stats))
    no_table_borders(stat_tbl)
    col_w = usable_cm / len(stats)
    for j, (num, label, color) in enumerate(stats):
        c = stat_tbl.rows[0].cells[j]
        set_cell_bg(c, color); no_cell_borders(c)
        cell_margin(c, top=70, bottom=70, left=60, right=60)
        stat_tbl.columns[j].width = Cm(col_w)
        p_n = c.paragraphs[0]
        p_n.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_n.paragraph_format.space_before = Pt(0)
        p_n.paragraph_format.space_after  = Pt(1)
        r_n = p_n.add_run(num)
        r_n.bold = True; r_n.font.size = Pt(19)
        r_n.font.color.rgb = _hex(C_WHITE)
        p_l = c.add_paragraph()
        p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_l.paragraph_format.space_before = Pt(0)
        p_l.paragraph_format.space_after  = Pt(0)
        r_l = p_l.add_run(label)
        r_l.font.size = Pt(6.5); r_l.bold = True
        r_l.font.color.rgb = _hex("d6eaf8")

    # ══════════════════════════════════════════════════════════════════════════
    # BLOCK 3 — PIPELINE FLOW DIAGRAM
    # ══════════════════════════════════════════════════════════════════════════
    _zero_para(doc)
    pipeline_steps = [
        ("DATA",       "Meteostat API",           "154360"),
        ("PREPROCESS", "Clean · Scale · IQR",     "1a5276"),
        ("MODELS",     "GRU·LSTM·RNN·LR·ARIMA",   "1f3864"),
        ("EVALUATE",   "MSE · MAE · R² · Acc",    "4a235a"),
        ("FORECAST",   "7-Day Rolling",            "145a32"),
    ]
    n_steps   = len(pipeline_steps)
    total_cols = n_steps * 2 - 1   # boxes + arrows

    pipe_tbl = doc.add_table(rows=1, cols=total_cols)
    no_table_borders(pipe_tbl)
    box_w   = (usable_cm * 0.86) / n_steps
    arrow_w = (usable_cm * 0.14) / (n_steps - 1)

    col_idx = 0
    for i, (label, sub, color) in enumerate(pipeline_steps):
        c = pipe_tbl.rows[0].cells[col_idx]
        set_cell_bg(c, color); no_cell_borders(c)
        cell_margin(c, top=50, bottom=50, left=40, right=40)
        pipe_tbl.columns[col_idx].width = Cm(box_w)
        p1 = c.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after  = Pt(1)
        r1 = p1.add_run(label)
        r1.bold = True; r1.font.size = Pt(7.5)
        r1.font.color.rgb = _hex(C_WHITE)
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        r2 = p2.add_run(sub)
        r2.font.size = Pt(6); r2.italic = True
        r2.font.color.rgb = _hex("adc8e0")
        col_idx += 1
        if i < n_steps - 1:
            ac = pipe_tbl.rows[0].cells[col_idx]
            set_cell_bg(ac, "d5d8dc"); no_cell_borders(ac)
            cell_margin(ac, top=50, bottom=50, left=5, right=5)
            pipe_tbl.columns[col_idx].width = Cm(arrow_w)
            pa = ac.paragraphs[0]
            pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pa.paragraph_format.space_before = Pt(10)
            pa.paragraph_format.space_after  = Pt(0)
            ra = pa.add_run("▶")
            ra.font.size = Pt(8); ra.font.color.rgb = _hex("7f8c8d")
            col_idx += 1

    # ══════════════════════════════════════════════════════════════════════════
    # BLOCK 4 — TWO-COLUMN BODY (no gap column — thin border divider instead)
    # ══════════════════════════════════════════════════════════════════════════
    _zero_para(doc)
    body_tbl = doc.add_table(rows=1, cols=2)
    no_table_borders(body_tbl)
    left_c  = body_tbl.rows[0].cells[0]
    right_c = body_tbl.rows[0].cells[1]
    L_W, R_W = 9.1, 9.49
    body_tbl.columns[0].width = Cm(L_W)
    body_tbl.columns[1].width = Cm(R_W)
    no_cell_borders(left_c); no_cell_borders(right_c)
    add_col_divider(left_c, color="bdc3c7", sz=6)   # thin gray line — no block of gray
    cell_margin(left_c,  top=50, bottom=50, left=50, right=140)
    cell_margin(right_c, top=50, bottom=50, left=140, right=50)

    # ── LEFT helpers ───────────────────────────────────────────────────────────
    def lsec(title, sb=6):
        p = left_c.add_paragraph()
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run("●  " + title.upper())
        r.bold = True; r.font.size = Pt(8.5)
        r.font.color.rgb = _hex(C_NAVY)

    def lbody(text, size=8.7, sa=2):
        p = left_c.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(sa)
        r = p.add_run(text); r.font.size = Pt(size)

    def lbul(text):
        p = left_c.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(text); r.font.size = Pt(8.5)

    # ── LEFT content ───────────────────────────────────────────────────────────
    lsec("Overview", sb=0)
    lbody(
        "A web-based benchmarking platform that fetches real daily weather data "
        "from Meteostat and evaluates five machine learning models for "
        "multi-variable weather forecasting across two Philippine stations.",
        sa=2,
    )
    lsec("Problem Statement")
    lbody(
        "Philippine weather forecasting relies on general-purpose models not "
        "specifically benchmarked on local station data. This study measures "
        "whether GRU outperforms alternative deep learning and statistical "
        "approaches on four meteorological variables.",
        sa=2,
    )
    lsec("Target Variables")
    for v in [
        "Precipitation (prcp)  —  daily rainfall, mm",
        "Temperature (tavg)     —  average daily temp, °C",
        "Wind Speed (wspd)     —  average wind speed, km/h",
        "Pressure (pres)          —  sea-level pressure, hPa",
    ]:
        lbul(v)
    lsec("Stations Covered")
    lbul("Baguio City  (Meteostat Station ID: 98328)")
    lbul("Manila           (Meteostat Station ID: 98425)")
    lsec("Data Period")
    lbody("Meteostat Open Weather API  ·  January 1, 2020 — Present (dynamic)")
    lsec("Methodology")
    lbody(
        "Data cleaned via IQR outlier removal and MinMax scaling. "
        "Deep learning models use sliding-window sequences (seq_len=30) with "
        "early stopping (patience=20). Results cached for instant reload.",
        sa=0,
    )

    # ── RIGHT helpers ──────────────────────────────────────────────────────────
    def rsec(title, sb=6):
        p = right_c.add_paragraph()
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run("●  " + title.upper())
        r.bold = True; r.font.size = Pt(8.5)
        r.font.color.rgb = _hex(C_NAVY)

    def rbul(text):
        p = right_c.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(text); r.font.size = Pt(8.5)

    # ── RIGHT: Model chips ─────────────────────────────────────────────────────
    rsec("Models Compared", sb=0)
    chips_tbl = right_c.add_table(rows=1, cols=len(MODEL_CHIPS))
    no_table_borders(chips_tbl)
    chip_w = R_W / len(MODEL_CHIPS)
    for j, (name, color) in enumerate(MODEL_CHIPS):
        cc = chips_tbl.rows[0].cells[j]
        set_cell_bg(cc, color); no_cell_borders(cc)
        cell_margin(cc, top=45, bottom=45, left=40, right=40)
        chips_tbl.columns[j].width = Cm(chip_w)
        pp2 = cc.paragraphs[0]
        pp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp2.paragraph_format.space_before = Pt(0)
        pp2.paragraph_format.space_after  = Pt(0)
        rr = pp2.add_run(name)
        rr.bold = True; rr.font.size = Pt(7.5)
        rr.font.color.rgb = _hex(C_WHITE)

    p_arch = right_c.add_paragraph()
    p_arch.paragraph_format.space_before = Pt(2)
    p_arch.paragraph_format.space_after  = Pt(2)
    r_arch = p_arch.add_run(
        "GRU & LSTM: 3-layer Bidirectional, hidden=256, dropout=0.2, epochs=200\n"
        "SimpleRNN: Unidirectional  ·  LR: flattened multivariate  ·  ARIMA: auto_arima (AIC)"
    )
    r_arch.font.size = Pt(7.5); r_arch.italic = True
    r_arch.font.color.rgb = _hex("555555")

    # ── RIGHT: Evaluation Metrics ──────────────────────────────────────────────
    rsec("Evaluation Metrics")
    for m in [
        "Accuracy (%)  =  100 − MAPE",
        "Mean Squared Error (MSE)",
        "Mean Absolute Error (MAE)",
        "Coefficient of Determination (R²)",
    ]:
        rbul(m)

    # ── RIGHT: Technology Stack ────────────────────────────────────────────────
    rsec("Technology Stack")
    stack = [
        ("Language",      "Python 3.10+"),
        ("Deep Learning", "PyTorch 2.x  (CUDA / MPS / CPU)"),
        ("ARIMA",         "pmdarima  auto_arima"),
        ("Data",          "Meteostat API"),
        ("Preprocessing", "scikit-learn  ·  pandas  ·  NumPy"),
        ("UI / Plots",    "Streamlit  ·  Matplotlib"),
    ]
    stbl = right_c.add_table(rows=len(stack), cols=2)
    no_table_borders(stbl)
    for i, (k, v) in enumerate(stack):
        c0, c1 = stbl.rows[i].cells
        bg = C_LIGHT if i % 2 == 0 else C_WHITE
        set_cell_bg(c0, bg); set_cell_bg(c1, bg)
        no_cell_borders(c0); no_cell_borders(c1)
        cell_margin(c0, top=28, bottom=28, left=80, right=40)
        cell_margin(c1, top=28, bottom=28, left=40, right=80)
        stbl.columns[0].width = Cm(3.0)
        stbl.columns[1].width = Cm(6.49)
        r0 = c0.paragraphs[0].add_run(k)
        r0.bold = True; r0.font.size = Pt(8)
        r1 = c1.paragraphs[0].add_run(v)
        r1.font.size = Pt(8)

    # ── RIGHT: Links ───────────────────────────────────────────────────────────
    rsec("Live Demo & Source Code")
    link_data = [
        ("\U0001f310  Live App",
         "thesis-gru-evaluation-in-weather-forecasting.streamlit.app",
         "1f3864", C_LIGHT),
        ("\U0001f4bb  Source Code",
         "github.com/EzekelCastro/Thesis-GRU-Evaluation-in-Weather-Forecasting",
         "2c3e50", "f4f6f7"),
    ]
    link_tbl = right_c.add_table(rows=len(link_data), cols=1)
    no_table_borders(link_tbl)
    for i, (label, url, color, bg) in enumerate(link_data):
        lc = link_tbl.rows[i].cells[0]
        set_cell_bg(lc, bg); no_cell_borders(lc)
        cell_margin(lc, top=35, bottom=35, left=80, right=80)
        p_lbl = lc.paragraphs[0]
        p_lbl.paragraph_format.space_before = Pt(0)
        p_lbl.paragraph_format.space_after  = Pt(0)
        r_lbl = p_lbl.add_run(label + "  ")
        r_lbl.bold = True; r_lbl.font.size = Pt(8)
        r_lbl.font.color.rgb = _hex(color)
        r_url = p_lbl.add_run(url)
        r_url.font.size = Pt(7.5); r_url.underline = True
        r_url.font.color.rgb = _hex("2471a3")

    # ══════════════════════════════════════════════════════════════════════════
    # BLOCK 5 — FOOTER BAR
    # ══════════════════════════════════════════════════════════════════════════
    _zero_para(doc)
    foot_tbl = doc.add_table(rows=1, cols=1)
    no_table_borders(foot_tbl)
    foot_c = foot_tbl.rows[0].cells[0]
    set_cell_bg(foot_c, C_DARK); no_cell_borders(foot_c)
    cell_margin(foot_c, top=55, bottom=55, left=200, right=200)
    pf = foot_c.paragraphs[0]
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.space_before = Pt(0)
    pf.paragraph_format.space_after  = Pt(0)
    rf = pf.add_run(
        f"University of the Cordilleras  ·  Baguio City, Benguet, Philippines  ·  {datetime.date.today().year}"
    )
    rf.font.size = Pt(7.5); rf.font.color.rgb = _hex(C_STEEL)

    doc.save("tech_brief.docx")
    print("Saved: tech_brief.docx")


# ══════════════════════════════════════════════════════════════════════════════
# 2. VALIDATION QUESTIONNAIRE
# ══════════════════════════════════════════════════════════════════════════════

def build_validation():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title block ────────────────────────────────────────────────────────────
    title_tbl = doc.add_table(rows=1, cols=1)
    no_table_borders(title_tbl)
    tc = title_tbl.rows[0].cells[0]
    set_cell_bg(tc, C_NAVY)
    no_cell_borders(tc)
    cell_margin(tc, top=140, bottom=140, left=200, right=200)

    p1 = tc.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after  = Pt(4)
    r1 = p1.add_run("SYSTEM VALIDATION QUESTIONNAIRE")
    r1.bold = True; r1.font.size = Pt(13)
    r1.font.color.rgb = _hex(C_WHITE)

    p2 = tc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(0)
    r2 = p2.add_run(
        "GRU Evaluation in Multi-Parameter Weather Forecasting System\n"
        "University of the Cordilleras  ·  B.S. Computer Science"
    )
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = _hex("adc8e0")

    # ── Instructions ───────────────────────────────────────────────────────────
    inst = doc.add_paragraph()
    inst.paragraph_format.space_before = Pt(10)
    inst.paragraph_format.space_after  = Pt(6)
    ir = inst.add_run(
        "Instructions:  Please rate each statement using the scale below. "
        "Write the number that best reflects your evaluation in the 'Rating' column. "
        "Additional comments are welcome in the 'Remarks' column."
    )
    ir.font.size = Pt(9.5)
    ir.italic = True

    # ── Scale legend ───────────────────────────────────────────────────────────
    scale_tbl = doc.add_table(rows=1, cols=5)
    no_table_borders(scale_tbl)
    headers = ["5 — Strongly Agree", "4 — Agree", "3 — Neutral",
               "2 — Disagree", "1 — Strongly Disagree"]
    colors  = ["1f3864", "2980b9", "7f8c8d", "e67e22", "c0392b"]
    for j, (h, c) in enumerate(zip(headers, colors)):
        cell = scale_tbl.rows[0].cells[j]
        set_cell_bg(cell, c)
        no_cell_borders(cell)
        cell_margin(cell, top=60, bottom=60, left=60, right=60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(8.5)
        r.font.color.rgb = _hex(C_WHITE)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Question sections ──────────────────────────────────────────────────────
    sections_data = [
        (
            "A.  Usability",
            "1f3864",
            [
                "The system interface is clean and easy to navigate.",
                "The sidebar controls (station, model selection, settings) are intuitive and easy to understand.",
                "The labels, axis titles, and chart legends are clear and informative.",
                "The system provides sufficient feedback during long-running operations (e.g., progress bar, status messages).",
                "I was able to run the analysis and interpret the results without needing external assistance.",
                "The tabs (Predictions, Metrics, Forecast, Download) are logically organized and easy to switch between.",
                "The 7-day Forecast tab is useful and easy to operate.",
                "The download feature works as expected and produces a usable output.",
            ],
        ),
        (
            "B.  System Performance",
            "1a5276",
            [
                "The system loads and responds within a reasonable time frame.",
                "The progress indicator accurately reflects the training progress.",
                "The system handles multiple models and stations without crashing or freezing.",
                "The caching mechanism effectively reduces wait time on subsequent loads.",
                "The system works consistently across multiple runs with the same settings.",
                "Error messages or warnings, if any, are understandable and helpful.",
            ],
        ),
        (
            "C.  Accuracy & Reliability of Results",
            "145a32",
            [
                "The prediction plots clearly show the difference between actual and predicted values.",
                "The metrics (MSE, MAE, R², Accuracy) are presented in a way that allows meaningful comparison across models.",
                "The bar charts in the Metrics tab effectively visualize model performance differences.",
                "The 'Best Model per Variable' summary table is helpful for drawing conclusions.",
                "The 7-day forecast values appear reasonable and consistent with historical trends.",
                "The results appear consistent and trustworthy based on the data presented.",
            ],
        ),
        (
            "D.  Relevance & Completeness",
            "4a235a",
            [
                "The choice of weather variables (Precipitation, Temperature, Wind Speed, Pressure) is appropriate for the study.",
                "The selection of stations (Baguio City and Manila) is relevant to the Philippine context.",
                "The five models compared (GRU, LSTM, SimpleRNN, Linear Regression, ARIMA) represent a sufficient range for benchmarking.",
                "The evaluation metrics used are appropriate for assessing forecasting model performance.",
                "The system as a whole addresses the research objectives of the study.",
            ],
        ),
        (
            "E.  Overall Assessment",
            "2c3e50",
            [
                "The system is suitable as a tool for evaluating weather forecasting models in an academic context.",
                "I would recommend this system as a reference for similar comparative studies.",
                "Overall, I am satisfied with the quality and functionality of this system.",
            ],
        ),
    ]

    for sec_title, color_hex, questions in sections_data:
        sh = doc.add_paragraph()
        sh.paragraph_format.space_before = Pt(10)
        sh.paragraph_format.space_after  = Pt(4)
        r = sh.add_run(sec_title)
        r.bold = True; r.font.size = Pt(11)
        c = bytes.fromhex(color_hex)
        r.font.color.rgb = RGBColor(c[0], c[1], c[2])

        qtbl = doc.add_table(rows=1 + len(questions), cols=3)
        qtbl.style = "Table Grid"
        qtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        qtbl.columns[0].width = Cm(0.8)
        qtbl.columns[1].width = Cm(10.5)
        qtbl.columns[2].width = Cm(2.5)

        hdr_cells = qtbl.rows[0].cells
        for j, lbl in enumerate(["No.", "Statement", "Rating (1–5)"]):
            set_cell_bg(hdr_cells[j], color_hex)
            p = hdr_cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(lbl)
            r.bold = True; r.font.size = Pt(9)
            r.font.color.rgb = _hex(C_WHITE)

        for i, q in enumerate(questions):
            row = qtbl.rows[i + 1]
            bg = "eaf0fb" if i % 2 == 0 else C_WHITE
            set_cell_bg(row.cells[0], bg)
            p0 = row.cells[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r0 = p0.add_run(str(i + 1))
            r0.font.size = Pt(9)
            set_cell_bg(row.cells[1], bg)
            r1 = row.cells[1].paragraphs[0].add_run(q)
            r1.font.size = Pt(9.5)
            set_cell_bg(row.cells[2], bg)
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Open-ended section ─────────────────────────────────────────────────────
    divider(doc)
    oe = doc.add_paragraph()
    oe.paragraph_format.space_before = Pt(8)
    oer = oe.add_run("F.  Open-Ended Feedback")
    oer.bold = True; oer.font.size = Pt(11)
    oer.font.color.rgb = _hex("2c3e50")

    open_qs = [
        "What aspect of the system do you find most useful or well-implemented?",
        "What aspect of the system needs the most improvement?",
        "Do you have any suggestions for additional features or changes to the current design?",
        "Any other comments, observations, or recommendations?",
    ]
    for q in open_qs:
        pq = doc.add_paragraph()
        pq.paragraph_format.space_before = Pt(8)
        pq.paragraph_format.space_after  = Pt(0)
        rq = pq.add_run(q)
        rq.bold = True; rq.font.size = Pt(9.5)
        for _ in range(2):
            pl = doc.add_paragraph()
            pl.paragraph_format.space_before = Pt(2)
            pl.paragraph_format.space_after  = Pt(0)
            pPr = pl._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "aaaaaa")
            pBdr.append(bottom)
            pPr.append(pBdr)

    # ── Signature block ────────────────────────────────────────────────────────
    divider(doc)
    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(10)
    sig.paragraph_format.space_after  = Pt(0)
    sig.add_run("Validator's Name: ").bold = True
    sig.runs[0].font.size = Pt(9.5)
    sig.add_run("_" * 35 + "      ")
    sig.add_run("Signature: ").bold = True
    sig.runs[-1].font.size = Pt(9.5)
    sig.add_run("_" * 25)

    sig2 = doc.add_paragraph()
    sig2.paragraph_format.space_before = Pt(6)
    sig2.add_run("Position / Expertise: ").bold = True
    sig2.runs[0].font.size = Pt(9.5)
    sig2.add_run("_" * 30 + "      ")
    sig2.add_run("Date: ").bold = True
    sig2.runs[-1].font.size = Pt(9.5)
    sig2.add_run("_" * 20)

    doc.save("validation_questions.docx")
    print("Saved: validation_questions.docx")


if __name__ == "__main__":
    build_tech_brief()
    build_validation()
    print("\nDone.")
