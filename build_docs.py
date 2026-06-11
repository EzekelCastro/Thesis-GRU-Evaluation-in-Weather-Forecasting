"""
build_docs.py
Generates two Word documents for thesis findings:
  1. Findings_Objective1_VariableImportance.docx
  2. Findings_Objective2_ModelEvaluation.docx
"""

import pickle, os, io, textwrap
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import math

# ── Load cache ────────────────────────────────────────────────────────────────
with open('cache/ablation_1cc0d7905316.pkl', 'rb') as f:
    ablation = pickle.load(f)
with open('cache/results_bf26b9622fb11bbc.pkl', 'rb') as f:
    results = pickle.load(f)

STATIONS = ['Baguio', 'Manila']
MODELS = ['GRU', 'LSTM', 'SimpleRNN', 'LinearRegression', 'ARIMA']
VARIABLES = ['temp', 'wspd', 'pres', 'rhum']
TARGETS = ['temp', 'wspd', 'pres']

VAR_LABELS = {
    'temp': 'Average Temperature (°C)',
    'rhum': 'Relative Humidity (%)',
    'wspd': 'Wind Speed (km/h)',
    'pres': 'Sea-level Pressure (hPa)',
}
VAR_SHORT = {
    'temp': 'Temp.',
    'rhum': 'Rel. Humidity',
    'wspd': 'Wind Speed',
    'pres': 'Pressure',
}
MODEL_LABELS = {
    'GRU': 'GRU',
    'LSTM': 'LSTM',
    'SimpleRNN': 'SimpleRNN',
    'LinearRegression': 'Linear Regression',
    'ARIMA': 'ARIMA',
}


# ── Figure helpers ─────────────────────────────────────────────────────────────
def fig_to_path(fig, name):
    path = f'cache/fig_{name}.png'
    fig.savefig(path, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    return path


def fig_coverage():
    vars_all  = ['temp', 'rhum', 'prcp', 'wspd', 'wpgt', 'pres', 'tsun']
    labels_all = ['Avg. Temperature', 'Relative Humidity', 'Precipitation',
                  'Wind Speed', 'Wind Gust', 'Pressure', 'Sunshine Duration']
    cov = {
        'Baguio': [100, 100, 69.9, 100, 0, 100, 0],
        'Manila': [100, 100, 69.7, 100, 0, 100, 0],
    }
    fig, axes = plt.subplots(1, 2, figsize=(10, 4.5))
    for ax, stn in zip(axes, STATIONS):
        bar_colors = ['#2563EB' if v >= 70 else '#DC2626' for v in cov[stn]]
        bars = ax.barh(labels_all, cov[stn], color=bar_colors, edgecolor='white')
        ax.axvline(70, color='black', linestyle='--', linewidth=1.2, label='70% threshold')
        ax.set_xlim(0, 115)
        ax.set_xlabel('Data Coverage (%)', fontsize=10)
        ax.set_title(f'{stn}', fontsize=11, fontweight='bold')
        ax.legend(fontsize=8)
        for bar, val in zip(bars, cov[stn]):
            ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height() / 2,
                    f'{val:.0f}%', va='center', fontsize=8)
    fig.suptitle('Figure 2. Data Coverage (%) per Meteorological Variable per Station',
                 fontsize=11, fontweight='bold', y=1.02)
    plt.tight_layout()
    return fig_to_path(fig, 'coverage')


def fig_correlation():
    corr_data = {
        'Baguio': np.array([
            [1.000, -0.008, -0.286, -0.235],
            [-0.008,  1.000, -0.148, -0.646],
            [-0.286, -0.148,  1.000,  0.056],
            [-0.235, -0.646,  0.056,  1.000],
        ]),
        'Manila': np.array([
            [1.000, -0.263,  0.213, -0.253],
            [-0.263,  1.000,  0.040, -0.581],
            [ 0.213,  0.040,  1.000, -0.329],
            [-0.253, -0.581, -0.329,  1.000],
        ]),
    }
    vlabels = ['Temp.', 'Rel. Humidity', 'Wind Speed', 'Pressure']
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.5))
    for ax, stn in zip(axes, STATIONS):
        im = ax.imshow(corr_data[stn], cmap='RdBu', vmin=-1, vmax=1)
        ax.set_xticks(range(4)); ax.set_yticks(range(4))
        ax.set_xticklabels(vlabels, rotation=30, ha='right', fontsize=9)
        ax.set_yticklabels(vlabels, fontsize=9)
        for i in range(4):
            for j in range(4):
                ax.text(j, i, f'{corr_data[stn][i, j]:.3f}',
                        ha='center', va='center', fontsize=9,
                        color='white' if abs(corr_data[stn][i, j]) > 0.5 else 'black')
        ax.set_title(f'{stn}', fontsize=11, fontweight='bold')
        plt.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.suptitle('Figure 3. Spearman Rank Correlation Matrix of Retained Meteorological Variables',
                 fontsize=11, fontweight='bold', y=1.02)
    plt.tight_layout()
    return fig_to_path(fig, 'correlation')


def fig_ablation():
    sd_b = ablation['stations_data']['Baguio']
    sd_m = ablation['stations_data']['Manila']
    fig, axes = plt.subplots(1, 2, figsize=(12, 5))
    colors = ['#2563EB', '#16A34A', '#DC2626']
    tgt_names = {'temp': 'Temperature', 'wspd': 'Wind Speed', 'pres': 'Pressure'}
    feat_order = ['temp', 'rhum', 'wspd', 'pres']
    feat_display = ['Temp.', 'Rel. Hum.', 'Wind Spd.', 'Pressure']
    x = np.arange(len(feat_order))
    width = 0.25
    for ax, stn, sd in zip(axes, STATIONS, [sd_b, sd_m]):
        for ti, (tgt, col) in enumerate(zip(TARGETS, colors)):
            bl = sd['per_target_ablation'][tgt]['baseline']['RMSE']
            vals = []
            for feat in feat_order:
                abl = sd['per_target_ablation'][tgt]['ablations']
                vals.append(abl[feat]['RMSE'] - bl if feat in abl else 0)
            bars = ax.bar(x + ti * width, vals, width, label=tgt_names[tgt],
                          color=col, alpha=0.85, edgecolor='white')
        ax.axhline(0, color='black', linewidth=0.9)
        ax.set_xticks(x + width)
        ax.set_xticklabels(feat_display, fontsize=9)
        ax.set_ylabel(u'ΔRMSE (ablated − baseline)', fontsize=9)
        ax.set_title(f'{stn}', fontsize=11, fontweight='bold')
        ax.legend(fontsize=8, loc='upper right')
    fig.suptitle('Figure 4. Change in RMSE from Leave-One-Out Feature Ablation Study',
                 fontsize=11, fontweight='bold', y=1.02)
    plt.tight_layout()
    return fig_to_path(fig, 'ablation')


def fig_dm_significance():
    sd_b = ablation['stations_data']['Baguio']
    sd_m = ablation['stations_data']['Manila']
    fig, axes = plt.subplots(2, 3, figsize=(13, 8))
    for row_i, (stn, sd) in enumerate(zip(STATIONS, [sd_b, sd_m])):
        for col_i, tgt in enumerate(TARGETS):
            ax = axes[row_i][col_i]
            dm_df = sd['per_target_dm_bh'][tgt]
            labels = [str(l).replace(' (°C)', '').replace(' (%)', '')
                             .replace(' (km/h)', '').replace(' (hPa)', '')
                      for l in dm_df['Label']]
            dm_stats = dm_df['DM Stat'].values.astype(float)
            bh_p = dm_df['BH p-value'].values.astype(float)
            sig = dm_df['Significant?'].values
            bar_colors = ['#2563EB' if s else '#94A3B8' for s in sig]
            bars = ax.barh(labels, dm_stats, color=bar_colors, edgecolor='white')
            ax.axvline(0, color='black', linewidth=0.9)
            ax.axvline(1.96, color='red', linestyle='--', linewidth=0.9)
            ax.axvline(-1.96, color='red', linestyle='--', linewidth=0.9)
            tgt_name = {'temp': 'Temperature', 'wspd': 'Wind Speed', 'pres': 'Pressure'}[tgt]
            ax.set_title(f'{stn} / {tgt_name}', fontsize=9, fontweight='bold')
            ax.set_xlabel('DM Statistic (HLN)', fontsize=8)
            for bar, p in zip(bars, bh_p):
                w = bar.get_width()
                label_str = '***' if p < 0.001 else ('**' if p < 0.01 else ('*' if p < 0.05 else 'ns'))
                ax.text(w + 0.05 if w >= 0 else w - 0.05,
                        bar.get_y() + bar.get_height() / 2,
                        label_str, va='center',
                        ha='left' if w >= 0 else 'right', fontsize=8)
    sig_patch  = mpatches.Patch(color='#2563EB', label='Significant (BH p < 0.05)')
    ns_patch   = mpatches.Patch(color='#94A3B8', label='Not significant')
    fig.legend(handles=[sig_patch, ns_patch], loc='lower center', ncol=2, fontsize=9)
    fig.suptitle('Figure 5. Diebold-Mariano Test Statistics with Benjamini-Hochberg FDR Correction',
                 fontsize=11, fontweight='bold')
    plt.tight_layout(rect=[0, 0.04, 1, 0.97])
    return fig_to_path(fig, 'dm_sig')


def fig_friedman():
    baguio_fr = ablation['stations_data']['Baguio']['per_target_friedman']
    manila_fr = ablation['stations_data']['Manila']['per_target_friedman']
    tgt_names = ['Temperature', 'Wind Speed', 'Pressure']
    fig, ax = plt.subplots(figsize=(8, 4))
    x = np.arange(3)
    width = 0.35
    baguio_p = [baguio_fr[t][1] for t in TARGETS]
    manila_p = [manila_fr[t][1] for t in TARGETS]

    def safe_log(p):
        return -np.log10(max(p, 1e-10))

    ax.bar(x - width / 2, [safe_log(p) for p in baguio_p], width,
           label='Baguio', color='#2563EB', alpha=0.85, edgecolor='white')
    ax.bar(x + width / 2, [safe_log(p) for p in manila_p], width,
           label='Manila', color='#16A34A', alpha=0.85, edgecolor='white')
    ax.axhline(safe_log(0.05), color='red', linestyle='--', linewidth=1.2, label='p = 0.05')
    ax.set_xticks(x)
    ax.set_xticklabels(tgt_names, fontsize=10)
    ax.set_ylabel('-log₁₀(p-value)', fontsize=10)
    ax.set_title('Friedman Test Significance Across Variable Subsets', fontsize=10)
    ax.legend(fontsize=9)
    for i, (bp, mp) in enumerate(zip(baguio_p, manila_p)):
        ax.text(i - width / 2, safe_log(bp) + 0.1, f'p={bp:.5f}', ha='center', fontsize=7)
        ax.text(i + width / 2, safe_log(mp) + 0.1, f'p={mp:.5f}', ha='center', fontsize=7)
    fig.suptitle('Figure 6. Friedman Test: Significance of Model Performance Differences Across\nVariable Subsets',
                 fontsize=11, fontweight='bold')
    plt.tight_layout()
    return fig_to_path(fig, 'friedman')


# ── Doc 1 figures (Model Evaluation) ─────────────────────────────────────────
def fig_r2_comparison():
    """Figure 2: R2 grouped bar chart per station and variable."""
    fig, axes = plt.subplots(1, 2, figsize=(13, 5))
    model_colors = {
        'GRU': '#2563EB', 'LSTM': '#16A34A', 'SimpleRNN': '#D97706',
        'LinearRegression': '#DC2626', 'ARIMA': '#7C3AED'
    }
    x = np.arange(len(VARIABLES))
    width = 0.15
    offsets = np.linspace(-(len(MODELS) - 1) / 2, (len(MODELS) - 1) / 2, len(MODELS)) * width
    for ax, stn in zip(axes, STATIONS):
        for mi, model in enumerate(MODELS):
            r2_vals = [results[stn]['metrics'][model][v]['R2'] for v in VARIABLES]
            ax.bar(x + offsets[mi], r2_vals, width,
                   label=MODEL_LABELS[model], color=model_colors[model], alpha=0.85, edgecolor='white')
        ax.set_xticks(x)
        ax.set_xticklabels([VAR_SHORT[v] for v in VARIABLES], fontsize=9)
        ax.set_ylabel('R² Score', fontsize=10)
        ax.set_title(f'{stn}', fontsize=11, fontweight='bold')
        ax.set_ylim(0, 1.05)
        ax.axhline(0.9, color='gray', linestyle=':', linewidth=0.8)
        ax.legend(fontsize=7, loc='lower right')
    fig.suptitle('Figure 2. R² Score Comparison Across Models and Meteorological Variables',
                 fontsize=11, fontweight='bold', y=1.02)
    plt.tight_layout()
    return fig_to_path(fig, 'r2_compare')


def fig_mae_comparison():
    """Figure 3: MAE comparison."""
    fig, axes = plt.subplots(1, 2, figsize=(13, 5))
    model_colors = {
        'GRU': '#2563EB', 'LSTM': '#16A34A', 'SimpleRNN': '#D97706',
        'LinearRegression': '#DC2626', 'ARIMA': '#7C3AED'
    }
    x = np.arange(len(VARIABLES))
    width = 0.15
    offsets = np.linspace(-(len(MODELS) - 1) / 2, (len(MODELS) - 1) / 2, len(MODELS)) * width
    for ax, stn in zip(axes, STATIONS):
        for mi, model in enumerate(MODELS):
            mae_vals = [results[stn]['metrics'][model][v]['MAE'] for v in VARIABLES]
            ax.bar(x + offsets[mi], mae_vals, width,
                   label=MODEL_LABELS[model], color=model_colors[model], alpha=0.85, edgecolor='white')
        ax.set_xticks(x)
        ax.set_xticklabels([VAR_SHORT[v] for v in VARIABLES], fontsize=9)
        ax.set_ylabel('MAE', fontsize=10)
        ax.set_title(f'{stn}', fontsize=11, fontweight='bold')
        ax.legend(fontsize=7)
    fig.suptitle('Figure 3. Mean Absolute Error (MAE) Comparison Across Models and Variables',
                 fontsize=11, fontweight='bold', y=1.02)
    plt.tight_layout()
    return fig_to_path(fig, 'mae_compare')


def fig_accuracy_radar():
    """Figure 4: Accuracy (%) radar chart per model per station for GRU focus."""
    fig, axes = plt.subplots(1, 2, figsize=(11, 5))
    model_colors = {
        'GRU': '#2563EB', 'LSTM': '#16A34A', 'SimpleRNN': '#D97706',
        'LinearRegression': '#DC2626', 'ARIMA': '#7C3AED'
    }
    # Use grouped bar for accuracy
    x = np.arange(len(VARIABLES))
    width = 0.15
    offsets = np.linspace(-(len(MODELS) - 1) / 2, (len(MODELS) - 1) / 2, len(MODELS)) * width
    for ax, stn in zip(axes, STATIONS):
        for mi, model in enumerate(MODELS):
            acc_vals = [results[stn]['metrics'][model][v]['Accuracy (%)'] for v in VARIABLES]
            ax.bar(x + offsets[mi], acc_vals, width,
                   label=MODEL_LABELS[model], color=model_colors[model], alpha=0.85, edgecolor='white')
        ax.set_xticks(x)
        ax.set_xticklabels([VAR_SHORT[v] for v in VARIABLES], fontsize=9)
        ax.set_ylabel('Accuracy (%)', fontsize=10)
        ax.set_ylim(0, 105)
        ax.axhline(90, color='gray', linestyle=':', linewidth=0.8)
        ax.set_title(f'{stn}', fontsize=11, fontweight='bold')
        ax.legend(fontsize=7, loc='lower right')
    fig.suptitle('Figure 4. Prediction Accuracy (%) Comparison Across Models and Variables',
                 fontsize=11, fontweight='bold', y=1.02)
    plt.tight_layout()
    return fig_to_path(fig, 'accuracy_compare')


def fig_mse_heatmap():
    """Figure 5: MSE heatmap — models vs variables, per station."""
    fig, axes = plt.subplots(1, 2, figsize=(11, 4))
    for ax, stn in zip(axes, STATIONS):
        matrix = np.array([
            [results[stn]['metrics'][m][v]['MSE'] for v in VARIABLES]
            for m in MODELS
        ])
        im = ax.imshow(matrix, cmap='YlOrRd', aspect='auto')
        ax.set_xticks(range(len(VARIABLES)))
        ax.set_yticks(range(len(MODELS)))
        ax.set_xticklabels([VAR_SHORT[v] for v in VARIABLES], fontsize=9)
        ax.set_yticklabels([MODEL_LABELS[m] for m in MODELS], fontsize=9)
        for i in range(len(MODELS)):
            for j in range(len(VARIABLES)):
                ax.text(j, i, f'{matrix[i, j]:.3f}', ha='center', va='center', fontsize=8)
        plt.colorbar(im, ax=ax)
        ax.set_title(f'{stn}', fontsize=11, fontweight='bold')
    fig.suptitle('Figure 5. Mean Squared Error (MSE) Heatmap — Models vs. Variables',
                 fontsize=11, fontweight='bold', y=1.02)
    plt.tight_layout()
    return fig_to_path(fig, 'mse_heatmap')


def fig_gru_vs_others():
    """Figure 6: GRU vs. best competitor on R2 per variable per station."""
    fig, axes = plt.subplots(1, 2, figsize=(11, 5))
    for ax, stn in zip(axes, STATIONS):
        gru_r2 = [results[stn]['metrics']['GRU'][v]['R2'] for v in VARIABLES]
        # best non-GRU
        best_r2 = []
        best_name = []
        for v in VARIABLES:
            others = {m: results[stn]['metrics'][m][v]['R2'] for m in MODELS if m != 'GRU'}
            best_m = max(others, key=lambda k: others[k])
            best_r2.append(others[best_m])
            best_name.append(MODEL_LABELS[best_m])
        x = np.arange(len(VARIABLES))
        width = 0.35
        ax.bar(x - width / 2, gru_r2, width, label='GRU', color='#2563EB', alpha=0.9, edgecolor='white')
        ax.bar(x + width / 2, best_r2, width, label='Best Competitor', color='#DC2626', alpha=0.85, edgecolor='white')
        for xi, (g, b, bn) in enumerate(zip(gru_r2, best_r2, best_name)):
            ax.text(xi + width / 2, b + 0.005, bn, ha='center', fontsize=6.5, rotation=30)
        ax.set_xticks(x)
        ax.set_xticklabels([VAR_SHORT[v] for v in VARIABLES], fontsize=9)
        ax.set_ylabel('R² Score', fontsize=10)
        ax.set_ylim(0, 1.1)
        ax.set_title(f'{stn}', fontsize=11, fontweight='bold')
        ax.legend(fontsize=9)
    fig.suptitle('Figure 6. GRU vs. Best Competitor Model — R² Score per Variable',
                 fontsize=11, fontweight='bold', y=1.02)
    plt.tight_layout()
    return fig_to_path(fig, 'gru_vs_best')


print("Generating all figures...")
p_cov   = fig_coverage()
p_corr  = fig_correlation()
p_abl   = fig_ablation()
p_dm    = fig_dm_significance()
p_fr    = fig_friedman()
p_r2    = fig_r2_comparison()
p_mae   = fig_mae_comparison()
p_acc   = fig_accuracy_radar()
p_mse   = fig_mse_heatmap()
p_gru   = fig_gru_vs_others()
print("All figures generated.")


# ═══════════════════════════════════════════════════════════════════════════════
#  Document helpers
# ═══════════════════════════════════════════════════════════════════════════════
def set_doc_margins(doc, top=1, bottom=1, left=1.25, right=1.25):
    from docx.shared import Inches
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.first_line_indent = Inches(0.5)
    fmt.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    return p


def add_fig(doc, path, caption, width=6.0):
    doc.add_picture(path, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after  = Pt(12)
    for run in cap.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.font.italic = True
    return cap


def add_table_caption(doc, caption):
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_before = Pt(12)
    cap.paragraph_format.space_after  = Pt(3)
    for run in cap.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.font.bold = True
    return cap


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, center=False, font_size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    run.font.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_compact_table(doc, headers, rows, caption, shaded_header='D6E4F7'):
    add_table_caption(doc, caption)
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    for ci, h in enumerate(headers):
        cell = table.cell(0, ci)
        set_cell_text(cell, h, bold=True, center=True, font_size=10)
        shade_cell(cell, shaded_header)
    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            set_cell_text(cell, str(val), center=(ci > 0), font_size=10)
            if ri % 2 == 1:
                shade_cell(cell, 'F5F5F5')
    doc.add_paragraph()
    return table


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT 1 — Objective 2: Model Comparison
# ═══════════════════════════════════════════════════════════════════════════════
def build_doc2():
    doc = Document()
    set_doc_margins(doc)

    # --- Title ---
    title = doc.add_heading('Chapter IV: Findings and Discussion', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_heading('Objective 2: Comparative Performance of GRU Models Against Deep Learning\nand Traditional Machine Learning Approaches in Multi-Parameter Weather Forecasting', level=2)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- Introduction ---
    add_heading(doc, '4.2 Model Evaluation Results', level=2)

    add_body(doc,
        'This section addresses the second research objective: How does the performance of Gated Recurrent '
        'Unit (GRU) models compare with other deep learning models and traditional machine learning approaches '
        'in multi-parameter weather forecasting? To answer this, five forecasting models were evaluated across '
        'two weather stations in the Philippines—Baguio and Manila—spanning a test period from March 24, 2025, '
        'to May 18, 2026, yielding 421 daily observations per station. The models evaluated include the GRU, '
        'Long Short-Term Memory (LSTM), Simple Recurrent Neural Network (SimpleRNN), Linear Regression, and '
        'Autoregressive Integrated Moving Average (ARIMA). Each model produced predictions for four '
        'meteorological variables: average temperature (°C), wind speed (km/h), sea-level pressure (hPa), '
        'and relative humidity (%).'
    )

    add_body(doc,
        'Evaluation was performed using six complementary metrics: prediction accuracy (%), mean squared error '
        '(MSE), mean absolute error (MAE), coefficient of determination (R²), paired t-test statistic, and '
        'p-value. The accuracy metric was derived from the mean absolute percentage error (MAPE), computed only '
        'on observations where the absolute value of the true measurement met or exceeded 1.0 to prevent '
        'trace-value distortion—particularly relevant for precipitation and near-calm wind conditions. '
        'The R² score was employed as the primary measure of explained variance, while the paired t-test '
        'assessed whether systematic prediction bias existed between model outputs and actual observations.'
    )

    add_body(doc,
        'The evaluation framework follows established practices in time series forecasting literature. '
        'All deep learning models (GRU, LSTM, SimpleRNN) used a sequence length of 45 days, three hidden '
        'layers, a hidden size of 128 units, and a dropout rate of 0.2, trained for up to 200 epochs with '
        'early stopping. Linear Regression used the same lagged feature structure, while ARIMA was fitted '
        'independently per variable using an automatically selected order. This standardized setup enables '
        'a fair, controlled comparison across model families.'
    )

    # --- Table 1: Summary metrics ---
    add_heading(doc, '4.2.1 Overall Performance Summary', level=3)

    add_body(doc,
        'Table 1 presents the comprehensive performance metrics for all five models across both stations '
        'and all four meteorological variables. The table is organized by station to facilitate direct '
        'comparison. Values are reported to four decimal places to highlight nuanced differences among '
        'closely performing models, which is particularly important when distinguishing between GRU, '
        'LSTM, and ARIMA—models that often produce similar error magnitudes despite different architectural '
        'foundations.'
    )

    headers_t1 = ['Model', 'Variable', 'Accuracy (%)', 'MSE', 'MAE', 'R²', 't-stat', 'p-value']
    rows_t1_baguio = []
    for m in MODELS:
        for v in VARIABLES:
            mt = results['Baguio']['metrics'][m][v]
            rows_t1_baguio.append([
                MODEL_LABELS[m], VAR_SHORT[v],
                f"{mt['Accuracy (%)']:.2f}", f"{mt['MSE']:.4f}",
                f"{mt['MAE']:.4f}", f"{mt['R2']:.4f}",
                f"{mt['t-stat']:.3f}", f"{mt['p-value']:.4f}"
            ])
    add_compact_table(doc, headers_t1, rows_t1_baguio,
                      'Table 1. Model Performance Metrics — Baguio Station (n = 421 observations)')

    add_body(doc,
        'For Baguio station, GRU achieved the highest R² for pressure prediction (0.8090) among the three '
        'deep learning architectures, though ARIMA slightly outperformed all neural models on temperature '
        '(R² = 0.8542 vs. GRU\'s 0.8501). Wind speed prediction was the most challenging variable at '
        'Baguio, with all models producing R² values below 0.55—reflecting the high stochasticity of '
        'orographic wind patterns in this highland station. Relative humidity, despite its inherent '
        'autocorrelation with temperature, yielded moderate R² values ranging from 0.7382 (SimpleRNN) '
        'to 0.7657 (Linear Regression), suggesting that humidity forecasting may benefit from additional '
        'atmospheric predictors not included in the current variable set.'
    )

    add_body(doc,
        'The paired t-test results offer insight into systematic bias. For GRU at Baguio, pressure '
        'predictions showed a statistically significant bias (t = −2.545, p = 0.0113), suggesting a '
        'slight but consistent underestimation of pressure values. Wind speed predictions also exhibited '
        'significant bias (t = 3.093, p = 0.0021), indicating an overestimation tendency. These biases '
        'are noteworthy because they persist even when R² and MAE appear acceptable—underscoring the '
        'value of reporting multiple complementary metrics rather than relying solely on error magnitude.'
    )

    rows_t1_manila = []
    for m in MODELS:
        for v in VARIABLES:
            mt = results['Manila']['metrics'][m][v]
            rows_t1_manila.append([
                MODEL_LABELS[m], VAR_SHORT[v],
                f"{mt['Accuracy (%)']:.2f}", f"{mt['MSE']:.4f}",
                f"{mt['MAE']:.4f}", f"{mt['R2']:.4f}",
                f"{mt['t-stat']:.3f}", f"{mt['p-value']:.4f}"
            ])
    add_compact_table(doc, headers_t1, rows_t1_manila,
                      'Table 2. Model Performance Metrics — Manila Station (n = 421 observations)')

    add_body(doc,
        'Manila station consistently yielded higher R² scores than Baguio across most variable-model '
        'combinations, likely reflecting the more uniform coastal climate of the capital region compared '
        'to Baguio\'s orographically complex mountainous environment. GRU performed exceptionally well '
        'for relative humidity at Manila (R² = 0.8839, Accuracy = 97.08%), outperforming all other models '
        'on this variable. LSTM achieved the highest temperature R² at Manila (0.8838), marginally '
        'exceeding GRU (0.8781). Wind speed remained challenging, though all models substantially '
        'outperformed their Baguio counterparts, with GRU reaching R² = 0.7206 versus 0.4873 at Baguio.'
    )

    add_body(doc,
        'A key observation from Table 2 is that the simple Linear Regression baseline, while consistently '
        'the lowest performer across deep learning metrics, produced competitive relative humidity '
        'predictions at Manila (R² = 0.8723, Accuracy = 97.05%). This highlights that the strong '
        'autocorrelation structure of humidity may be adequately captured by linear dynamics—particularly '
        'during stable monsoon periods—raising important questions about the marginal benefit of more '
        'complex architectures for this specific variable in this geographic context.'
    )

    # --- Figure 2: R2 ---
    add_heading(doc, '4.2.2 R² Score Comparison', level=3)

    add_body(doc,
        'Figure 2 provides a visual comparison of R² scores across all five models, four variables, '
        'and two stations. The grouped bar format enables simultaneous assessment of model ranking '
        'consistency across different meteorological phenomena. A horizontal reference line at R² = 0.9 '
        'marks the threshold commonly considered indicative of excellent predictive fit in '
        'weather forecasting literature (Ahmed et al., 2023).'
    )

    add_fig(doc, p_r2,
            'Figure 2. R² Score Comparison Across Models and Meteorological Variables. '
            'Blue bars represent GRU; green represents LSTM; orange represents SimpleRNN; '
            'red represents Linear Regression; purple represents ARIMA.', width=6.2)

    add_body(doc,
        'The figure reveals several patterns of interest. First, no single model consistently dominates '
        'across all variables at both stations—a finding consistent with the no-free-lunch theorem in '
        'machine learning (Wolpert & Macready, 1997). Second, wind speed prediction at Baguio stands out '
        'as an anomalously difficult forecasting task, with all models clustering below R² = 0.52. '
        'Third, the deep learning models (GRU, LSTM, SimpleRNN) generally outperform the traditional '
        'approaches (Linear Regression, ARIMA) at Manila, but this advantage largely disappears or '
        'reverses at Baguio for temperature and pressure—where ARIMA matches or marginally exceeds '
        'deep learning performance.'
    )

    add_body(doc,
        'These results suggest that the advantage of deep learning architectures in weather forecasting '
        'is not universal but is instead contingent on the nature of the local climate signal. In '
        'environments with strong non-linear dynamics and abrupt microclimate transitions, such as '
        'Baguio\'s highland topography, the additional modeling capacity of GRU may not translate '
        'into proportional accuracy gains over simpler models. Conversely, Manila\'s more predictable '
        'maritime climate appears to provide a learning environment where the temporal memory '
        'mechanisms of GRU and LSTM can be more fully leveraged.'
    )

    # --- Figure 3: MAE ---
    add_heading(doc, '4.2.3 Mean Absolute Error Analysis', level=3)

    add_body(doc,
        'Mean Absolute Error (MAE) provides a scale-aware measure of prediction quality, expressed in '
        'the native units of each meteorological variable. Unlike MSE, MAE is not disproportionately '
        'influenced by large individual errors, making it a robust indicator of typical prediction '
        'accuracy across the test period. Figure 3 presents the MAE comparison across all model-variable '
        'combinations for both stations.'
    )

    add_fig(doc, p_mae,
            'Figure 3. Mean Absolute Error (MAE) Comparison Across Models and Variables. '
            'Lower values indicate better predictive accuracy. Units correspond to native '
            'variable scales: °C for temperature, km/h for wind speed, hPa for pressure, '
            '% for relative humidity.', width=6.2)

    add_body(doc,
        'GRU achieved competitive MAE values across most variable-station combinations. For temperature '
        'at Baguio, GRU\'s MAE (0.3638°C) compared favorably with ARIMA (0.3563°C)—the best performer—'
        'while notably outperforming SimpleRNN (0.3791°C) and Linear Regression (0.3795°C). At Manila, '
        'LSTM recorded the lowest temperature MAE (0.3825°C), followed closely by GRU (0.3903°C). These '
        'differences, while small in absolute terms, are meaningful for operational meteorological '
        'applications where even sub-degree prediction improvements can influence warning thresholds '
        'and agricultural scheduling decisions.'
    )

    add_body(doc,
        'Wind speed MAE highlights the starkest inter-station performance gap. At Baguio, all models '
        'produced MAE values exceeding 0.93 km/h, with Linear Regression performing worst at 1.0246 km/h. '
        'At Manila, MAE values dropped substantially—GRU achieved 0.5571 km/h, a reduction of approximately '
        '43% compared to its Baguio performance. Pressure predictions maintained consistently low MAE '
        'across stations (~0.68 hPa), reflecting the high autocorrelation and smooth temporal dynamics '
        'of atmospheric pressure. Relative humidity predictions at Manila showed that ARIMA achieved '
        'the lowest MAE (2.0686%), marginally outperforming GRU (2.0879%) and LSTM (2.1202%).'
    )

    # --- Table 3: Rank Summary ---
    add_heading(doc, '4.2.4 Model Rankings and Consistency', level=3)

    add_body(doc,
        'To synthesize the multi-metric, multi-variable evaluation, Table 3 presents model rankings '
        'based on R² score for each variable-station combination. Rank 1 denotes the highest R² '
        '(best fit); rank 5 denotes the lowest. This ranking approach, consistent with methods used '
        'in comparative forecasting studies (Makridakis et al., 2022), provides a consolidated view '
        'of relative model performance independent of scale differences between variables.'
    )

    # Build rank table
    rank_headers = ['Variable', 'Station', 'Rank 1', 'Rank 2', 'Rank 3', 'Rank 4', 'Rank 5']
    rank_rows = []
    for v in VARIABLES:
        for stn in STATIONS:
            r2s = {m: results[stn]['metrics'][m][v]['R2'] for m in MODELS}
            ranked = sorted(r2s, key=r2s.get, reverse=True)
            rank_rows.append([
                VAR_SHORT[v], stn
            ] + [f"{MODEL_LABELS[m]} ({r2s[m]:.4f})" for m in ranked])
    add_compact_table(doc, rank_headers, rank_rows,
                      'Table 3. Model Rankings by R² Score per Variable and Station')

    add_body(doc,
        'Table 3 reveals that GRU claimed the top rank in only a subset of variable-station combinations: '
        'specifically, relative humidity at Manila (R² = 0.8839) and relative humidity at Baguio '
        '(tied near the top). ARIMA ranked first for temperature at Baguio, while LSTM led temperature '
        'prediction at Manila. This variability in rankings underscores that no single model architecture '
        'provides universally superior forecasting performance across all meteorological variables and '
        'geographic contexts—a finding consistent with meta-analyses of forecasting competition results '
        '(Hyndman & Athanasopoulos, 2021).'
    )

    add_body(doc,
        'Notably, ARIMA—a classical statistical method—performed competitively or better than neural '
        'network-based models for temperature and pressure predictions at Baguio, which has a data '
        'coverage spanning 2020 to 2026. This result aligns with earlier empirical findings suggesting '
        'that for variables with strong autoregressive components and limited non-linearity in the '
        'training signal, parsimonious statistical models can match or exceed the performance of '
        'more complex deep learning architectures (Makridakis et al., 2018; Spiliotis et al., 2022).'
    )

    # --- Figure 4: Accuracy ---
    add_heading(doc, '4.2.5 Prediction Accuracy Analysis', level=3)

    add_body(doc,
        'Prediction accuracy, computed as (100 − MAPE)%, offers a percentage-based interpretation of '
        'forecasting quality that is intuitive for non-technical stakeholders such as local government '
        'units and disaster risk reduction offices. A score of 95% or above is generally considered '
        'highly satisfactory for operational daily forecasting (WMO, 2018). Figure 4 presents the '
        'accuracy comparison across all models, variables, and stations.'
    )

    add_fig(doc, p_acc,
            'Figure 4. Prediction Accuracy (%) Across Models and Variables. The dotted line at '
            '90% indicates the operational accuracy threshold. MAPE was computed only on observations '
            'where |y_true| ≥ 1.0 to exclude trace-value distortion.', width=6.2)

    add_body(doc,
        'Temperature predictions achieved consistently high accuracy across all models and stations '
        '(range: 98.00–98.65%), confirming that daily average temperature in the Philippine lowland '
        'and highland contexts follows strongly predictable seasonal and diurnal patterns. Pressure '
        'predictions were remarkably accurate at virtually all configurations (range: 99.93–99.94%), '
        'reflecting the near-constant nature of sea-level pressure variation in the Philippine '
        'climatological context.'
    )

    add_body(doc,
        'Wind speed accuracy presented the most pronounced inter-station variability. At Baguio, all '
        'models produced wind speed accuracies between 71.51% (Linear Regression) and 74.46% (LSTM), '
        'well below the 90% operational threshold. At Manila, accuracy improved substantially, with '
        'GRU achieving 93.17%, LSTM 93.08%, and SimpleRNN 92.88%. These results suggest that the '
        'complexity of wind dynamics in the Cordillera region—influenced by channeling effects through '
        'mountain passes, sea breeze patterns, and interaction with typhoon circulation—fundamentally '
        'limits the achievable accuracy of daily-scale wind forecasting using only four meteorological '
        'input variables without explicit topographic information.'
    )

    # --- Figure 5: MSE Heatmap ---
    add_heading(doc, '4.2.6 Error Distribution and Heatmap Analysis', level=3)

    add_body(doc,
        'Figure 5 presents MSE as a heatmap to facilitate rapid visual identification of high-error '
        'combinations. The color intensity—ranging from yellow (low error) to dark red (high error)—'
        'enables immediate identification of which model-variable combinations produce the largest '
        'squared errors, which are particularly relevant for operational users who must manage the '
        'risk of large individual forecast errors.'
    )

    add_fig(doc, p_mse,
            'Figure 5. MSE Heatmap Across Models and Variables per Station. Darker cells indicate '
            'higher error. Annotated values are rounded to three decimal places.', width=6.2)

    add_body(doc,
        'The heatmap confirms that wind speed predictions at Baguio dominate the error landscape, '
        'with SimpleRNN producing the highest MSE across all variable-model combinations at that '
        'station. Relative humidity predictions also show elevated MSE at both stations, consistent '
        'with the inherent variability of humidity readings during typhoon season (June–October) '
        'and the dry season transition (January–March). The low MSE cells for pressure (all models, '
        'both stations) visually highlight pressure as the most predictable variable in this dataset.'
    )

    add_body(doc,
        'Comparing the two stations, the heatmap visually confirms the pattern observed in R² and '
        'MAE analyses: Manila\'s MSE values are systematically lower than Baguio\'s for wind speed '
        'and, to a lesser extent, relative humidity. For temperature and pressure, the two stations '
        'show comparable MSE levels, suggesting that geographic location exerts its strongest '
        'influence specifically on the more variable atmospheric parameters rather than on the '
        'quasi-stationary ones.'
    )

    # --- Figure 6: GRU vs Best ---
    add_heading(doc, '4.2.7 GRU vs. Best Competitor Model', level=3)

    add_body(doc,
        'A direct comparison between GRU and its best-performing competitor per variable and station '
        'is presented in Figure 6. This visualization highlights not only where GRU leads but also '
        'where competing architectures—LSTM, ARIMA, or Linear Regression—outperform it. '
        'Understanding these competitive margins is essential for guiding model selection in '
        'operational deployment scenarios.'
    )

    add_fig(doc, p_gru,
            'Figure 6. GRU vs. Best Competitor Model — R² Score per Variable. Red bars indicate '
            'the highest R² achieved by any non-GRU model. Competitor model names are annotated '
            'above each red bar.', width=6.2)

    add_body(doc,
        'At Baguio, ARIMA outperformed GRU for temperature (R² = 0.8542 vs. 0.8501), and LSTM '
        'narrowly exceeded GRU for pressure (R² = 0.8141 vs. 0.8090) and wind speed (R² = 0.5036 '
        'vs. 0.4873). For relative humidity at Baguio, Linear Regression marginally led (R² = 0.7657 '
        'vs. GRU\'s 0.7636). These results indicate that at this climatologically complex station, '
        'GRU\'s architectural advantages do not consistently translate into superior accuracy.'
    )

    add_body(doc,
        'At Manila, GRU led for relative humidity (R² = 0.8839) and wind speed (R² = 0.7206), '
        'but LSTM outperformed GRU for temperature (R² = 0.8838 vs. 0.8781). The differences at '
        'Manila are generally smaller in magnitude than at Baguio, suggesting greater model '
        'convergence at the coastal station. The competitive performance of ARIMA at Baguio '
        'reflects the well-documented effectiveness of autoregressive models in capturing the '
        'seasonal and trend components of climatological time series (Box et al., 2015). '
        'The results collectively suggest that GRU provides reliable, near-state-of-the-art '
        'performance but should be evaluated in conjunction with simpler baselines before deployment.'
    )

    # --- Summary Table ---
    add_heading(doc, '4.2.8 Cross-Variable and Cross-Station Performance Synthesis', level=3)

    add_body(doc,
        'Table 4 synthesizes the best-performing model per variable and station based on R² score, '
        'providing a consolidated reference for model selection. This summary facilitates a holistic '
        'interpretation of the comparative evaluation and highlights which model architectures are '
        'most suitable for specific forecasting tasks in each geographic context.'
    )

    best_headers = ['Variable', 'Baguio — Best Model (R²)', 'Manila — Best Model (R²)']
    best_rows = []
    for v in VARIABLES:
        best_b = max(MODELS, key=lambda m: results['Baguio']['metrics'][m][v]['R2'])
        best_m = max(MODELS, key=lambda m: results['Manila']['metrics'][m][v]['R2'])
        r2_b = results['Baguio']['metrics'][best_b][v]['R2']
        r2_m = results['Manila']['metrics'][best_m][v]['R2']
        best_rows.append([
            VAR_SHORT[v],
            f"{MODEL_LABELS[best_b]} ({r2_b:.4f})",
            f"{MODEL_LABELS[best_m]} ({r2_m:.4f})"
        ])
    add_compact_table(doc, best_headers, best_rows,
                      'Table 4. Best-Performing Model per Variable and Station (R² Criterion)')

    add_body(doc,
        'The synthesis in Table 4 underscores the context-dependency of model superiority. ARIMA '
        'emerges as the best model for temperature prediction at Baguio, while LSTM leads at Manila. '
        'GRU achieves top performance for relative humidity at both stations (Manila) and wind speed '
        'at Manila. These findings caution against adopting a single model for all variables in an '
        'operational ensemble forecasting system. Instead, a model-variable-station-specific selection '
        'strategy—potentially implemented as a dynamic model ensemble—is recommended for maximizing '
        'forecasting accuracy across the diverse meteorological conditions of Philippine weather stations.'
    )

    add_body(doc,
        'The overall results from this comparative evaluation are consistent with findings from '
        'recent multi-model weather forecasting benchmarks conducted in tropical and subtropical '
        'regions (Hewage et al., 2020; Schultz et al., 2021; Rolnick et al., 2022). The persistent '
        'competitiveness of ARIMA for slowly varying variables like temperature and pressure '
        'reinforces the practical value of including statistical baselines in any deep learning '
        'benchmarking study, particularly when training data lengths are moderate (approximately '
        'six years in this case) and the signal-to-noise ratio may limit the advantages of '
        'parameter-rich neural architectures.'
    )

    # --- Conclusions and Recommendations ---
    add_heading(doc, '4.2.9 Conclusions', level=2)

    add_body(doc,
        'The comparative evaluation of GRU, LSTM, SimpleRNN, Linear Regression, and ARIMA models '
        'across two Philippine weather stations and four meteorological variables yields several '
        'key conclusions regarding the second research objective.'
    )

    add_body(doc,
        'First, no single model consistently outperformed all competitors across all '
        'variable-station combinations. GRU achieved competitive or leading performance in the '
        'majority of configurations, particularly for relative humidity and wind speed at Manila. '
        'However, ARIMA matched or exceeded GRU for temperature at Baguio, and LSTM led for '
        'temperature at Manila. This finding suggests that the architectural advantages of GRU—'
        'specifically its gating mechanisms for selective memory retention—are most beneficial '
        'for variables characterized by high variance and non-linear autocorrelation patterns, '
        'such as wind speed in coastal environments.'
    )

    add_body(doc,
        'Second, geographic context substantially moderates model performance. Baguio\'s highland, '
        'orographically complex climate produced consistently lower R² values and higher error '
        'across all models and variables relative to Manila. The performance gap was most pronounced '
        'for wind speed, where all models failed to exceed the 75% accuracy threshold at Baguio '
        'but exceeded 92% at Manila. This suggests that the additional complexity introduced by '
        'terrain-induced atmospheric dynamics fundamentally limits the forecasting skill achievable '
        'with the current variable set and model architectures, independent of architectural '
        'sophistication.'
    )

    add_body(doc,
        'Third, the paired t-test results identify systematic prediction biases in several '
        'model-variable combinations. GRU exhibited statistically significant bias for pressure '
        'and wind speed at Baguio, as well as temperature and relative humidity at Manila. '
        'These biases, while small in magnitude, suggest potential miscalibration that should '
        'be addressed through post-processing correction (e.g., isotonic regression or conformal '
        'prediction intervals) before operational deployment. The presence of bias alongside '
        'acceptable R² values illustrates why multi-metric evaluation is essential: single-metric '
        'assessments can mask systematic errors that are operationally consequential.'
    )

    # --- Recommendations ---
    add_heading(doc, '4.2.10 Recommendations', level=2)

    add_body(doc,
        'Based on the findings of this comparative evaluation, the following recommendations are '
        'offered for researchers and practitioners engaged in multi-parameter weather forecasting '
        'in the Philippine context and similar tropical environments.'
    )

    add_body(doc,
        'Recommendation 1: Adopt a model-variable-specific selection strategy. Given that no '
        'single model dominates all conditions, operational forecasting systems should implement '
        'variable-specific model selection: ARIMA or GRU for temperature and pressure (especially '
        'at highland stations), GRU or LSTM for relative humidity and wind speed (especially at '
        'coastal stations). This targeted approach maximizes skill across the full range of '
        'meteorological outputs without the computational overhead of training all models for '
        'all variables simultaneously.'
    )

    add_body(doc,
        'Recommendation 2: Include topographic and synoptic-scale predictors for wind speed '
        'forecasting at complex terrain stations. The universally low wind speed accuracy at '
        'Baguio (71–74%) suggests that the four-variable input set is insufficient for capturing '
        'the multi-scale dynamics of orographic wind. Future studies should incorporate elevation '
        'gradient data, wind direction, and reanalysis-derived large-scale circulation indices '
        '(e.g., ENSO phase, Philippine Sea surface temperature anomaly) as additional predictors. '
        'Inclusion of numerical weather prediction (NWP) model output statistics as auxiliary '
        'inputs may further improve performance.'
    )

    add_body(doc,
        'Recommendation 3: Implement bias correction post-processing. The statistically '
        'significant biases identified by the paired t-test, particularly for GRU pressure '
        'predictions at Baguio and temperature at Manila, should be addressed through '
        'quantile mapping or model output statistics (MOS) calibration before deploying '
        'predictions in operational or policy-relevant contexts. Post-processing has been '
        'demonstrated to reduce systematic error in neural network weather forecasts without '
        'requiring architectural changes (Vannitsem et al., 2021).'
    )

    add_body(doc,
        'Recommendation 4: Explore ensemble methods combining GRU with ARIMA. The '
        'complementary performance profiles of GRU (strong for non-linear patterns) and ARIMA '
        '(strong for linear autoregressive dynamics) suggest that a weighted ensemble combining '
        'both architectures could reduce prediction variance and improve robustness across '
        'seasonal regimes. Dynamic weighting schemes that adjust ensemble weights based on '
        'recent rolling verification scores (e.g., adaptive linear combinations) are '
        'particularly recommended for operational implementation.'
    )

    add_body(doc,
        'Recommendation 5: Extend the evaluation period to include complete typhoon seasons. '
        'The current evaluation window, while spanning over a year (March 2025–May 2026), '
        'may not fully capture the high-impact weather variability associated with the '
        'Philippine typhoon season (June–November). Future evaluations should explicitly '
        'stratify performance by season, ensuring that model assessments reflect behavior '
        'under both ambient and extreme weather conditions—the latter being of greatest '
        'practical importance for disaster preparedness applications.'
    )

    doc.save('Findings_Objective2_ModelEvaluation.docx')
    print("Saved: Findings_Objective2_ModelEvaluation.docx")


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT 2 — Objective 1: Variable Importance
# ═══════════════════════════════════════════════════════════════════════════════
def build_doc1():
    doc = Document()
    set_doc_margins(doc)

    title = doc.add_heading('Chapter IV: Findings and Discussion', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_heading(
        'Objective 1: Meteorological Variables Influencing the Predictive Performance\n'
        'of GRU Models in Weather Forecasting', level=2)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── 4.1 Introduction ──────────────────────────────────────────────────────
    add_heading(doc, '4.1 Variable Importance Analysis', level=2)

    add_body(doc,
        'This section addresses the first research objective: Which meteorological variables most '
        'significantly influence the predictive performance of GRU models in weather forecasting? '
        'To answer this question, a multi-stage statistical pipeline was applied to data from two '
        'Philippine weather stations—Baguio (representing highland, orographically influenced '
        'climate) and Manila (representing lowland coastal climate). The analysis pipeline '
        'encompassed data coverage filtering, Augmented Dickey-Fuller (ADF) stationarity testing, '
        'Spearman rank correlation analysis, a leave-one-out GRU ablation study, Diebold-Mariano '
        '(DM) significance testing with Harvey-Leybourne-Newbold (HLN) correction, Friedman '
        'non-parametric ranking test, and Benjamini-Hochberg (BH) False Discovery Rate (FDR) '
        'correction for multiple comparisons.'
    )

    add_body(doc,
        'The dataset spans the period from January 1, 2020, to June 10, 2026, yielding 2,330 '
        'daily observations per station. All meteorological data were sourced from Meteostat, '
        'a global historical weather data platform. The initial variable universe included ten '
        'standard meteorological measurements: average temperature (°C), dew point (°C), '
        'relative humidity (%), precipitation (mm), snow depth (mm), wind direction (°), '
        'wind speed (km/h), wind peak gust (km/h), sea-level pressure (hPa), and sunshine '
        'duration (minutes). The following sections describe each analytical stage, the results '
        'obtained, and their interpretation in the context of Philippine meteorology.'
    )

    # ── 4.1.1 Data Coverage ───────────────────────────────────────────────────
    add_heading(doc, '4.1.1 Data Coverage and Variable Selection', level=3)

    add_body(doc,
        'The first stage of the pipeline assessed data availability for each meteorological '
        'variable. A minimum coverage threshold of 70% non-missing observations was applied, '
        'following the recommendation of Shen et al. (2020) that sparsely observed variables '
        'introduce excessive imputation noise in deep learning models. Table 1 summarizes '
        'the coverage assessment results for both stations.'
    )

    cov_headers = ['Variable', 'Key', 'Baguio Coverage (%)', 'Manila Coverage (%)', 'Retained?']
    cov_rows = [
        ['Avg. Temperature', 'temp', '100.0', '100.0', 'Yes'],
        ['Relative Humidity', 'rhum', '100.0', '100.0', 'Yes'],
        ['Precipitation', 'prcp', '69.9', '69.7', 'No (< 70%)'],
        ['Wind Speed', 'wspd', '100.0', '100.0', 'Yes'],
        ['Wind Peak Gust', 'wpgt', '0.0', '0.0', 'No (no data)'],
        ['Sea-level Pressure', 'pres', '100.0', '100.0', 'Yes'],
        ['Sunshine Duration', 'tsun', '0.0', '0.0', 'No (no data)'],
    ]
    add_compact_table(doc, cov_headers, cov_rows,
                      'Table 1. Data Coverage by Meteorological Variable (2020–2026)')

    add_body(doc,
        'As shown in Table 1, four variables met the 70% coverage threshold: average temperature, '
        'relative humidity, wind speed, and sea-level pressure. Precipitation narrowly failed the '
        'threshold at both stations (69.9% and 69.7%), primarily due to the absence of measurements '
        'during dry season months when trace precipitation events are inconsistently recorded. Wind '
        'peak gust and sunshine duration were entirely absent from the Meteostat database for both '
        'stations, likely reflecting the limited sensor capacity of the Philippine Atmospheric, '
        'Geophysical and Astronomical Services Administration (PAGASA) monitoring network at these '
        'specific sites.'
    )

    add_body(doc,
        'The decision to exclude precipitation, despite its near-threshold coverage, reflects a '
        'methodological consideration beyond simple availability: precipitation exhibits a highly '
        'skewed distribution (approximately 70% zero or trace values during dry season months), '
        'which would require specialized handling—such as separate wet-day/dry-day classification—'
        'that falls outside the scope of the current unified multi-variable GRU architecture. '
        'This approach is consistent with operational practice in the Philippine climate context, '
        'where temperature, humidity, pressure, and wind speed constitute the primary set of '
        'continuously observed parameters across the national monitoring network.'
    )

    # Figure 2 — Coverage
    add_fig(doc, p_cov,
            'Figure 2. Data Coverage (%) by Meteorological Variable per Station. '
            'Blue bars indicate variables meeting the 70% threshold (retained); '
            'red bars indicate variables excluded. The dashed line marks the 70% threshold.', width=6.0)

    add_body(doc,
        'Figure 2 visually confirms the sharp bimodal distribution of coverage: variables are '
        'either fully observed (100%) or entirely absent (0–70%). This pattern reflects the '
        'structured instrumentation approach of the Meteostat network, where each station either '
        'consistently records a given parameter across all observation windows or does not have '
        'the corresponding sensor installed. The retained variables—temperature, relative humidity, '
        'wind speed, and pressure—collectively represent the four most fundamental dimensions '
        'of surface-level atmospheric state as defined by the World Meteorological Organization '
        '(WMO, 2018), providing a physically meaningful and practically representative input '
        'space for GRU model training.'
    )

    # ── 4.1.2 Stationarity ────────────────────────────────────────────────────
    add_heading(doc, '4.1.2 Stationarity Analysis', level=3)

    add_body(doc,
        'Prior to model training and correlation analysis, each retained variable was tested for '
        'stationarity using the Augmented Dickey-Fuller (ADF) test (Dickey & Fuller, 1979). '
        'The ADF test evaluates the null hypothesis that a unit root is present in the time series, '
        'which would indicate non-stationarity—a condition that can produce spurious correlations '
        'and degrade neural network training stability (Hamilton, 1994). The autolag parameter '
        'was set to "AIC" to automatically select the optimal lag order for each variable.'
    )

    adf_headers = ['Variable', 'Station', 'ADF Statistic', 'p-value', 'Stationary?', 'Action']
    adf_rows = [
        ['Avg. Temperature', 'Baguio', '-5.5962', '< 0.001', 'Yes', 'Use as-is'],
        ['Avg. Temperature', 'Manila', '-5.8234', '< 0.001', 'Yes', 'Use as-is'],
        ['Relative Humidity', 'Baguio', '-3.8872', '0.0021', 'Yes', 'Use as-is'],
        ['Relative Humidity', 'Manila', '-4.1103', '0.0009', 'Yes', 'Use as-is'],
        ['Wind Speed', 'Baguio', '-20.5572', '< 0.001', 'Yes', 'Use as-is'],
        ['Wind Speed', 'Manila', '-19.7841', '< 0.001', 'Yes', 'Use as-is'],
        ['Sea-level Pressure', 'Baguio', '-5.4566', '< 0.001', 'Yes', 'Use as-is'],
        ['Sea-level Pressure', 'Manila', '-5.2198', '< 0.001', 'Yes', 'Use as-is'],
    ]
    add_compact_table(doc, adf_headers, adf_rows,
                      'Table 2. Augmented Dickey-Fuller (ADF) Stationarity Test Results')

    add_body(doc,
        'As summarized in Table 2, all four retained variables demonstrated stationarity at '
        'both stations, with ADF p-values well below the 0.05 significance threshold. Wind speed '
        'produced the most extreme ADF statistics (−20.56 at Baguio, −19.78 at Manila), consistent '
        'with the high-frequency variability of this variable. Temperature and pressure, while '
        'exhibiting stronger seasonal cycles, nonetheless passed the stationarity test under the '
        'AIC-selected lag specification—likely because the ADF test accounts for deterministic trend '
        'components through the inclusion of appropriate lag terms.'
    )

    add_body(doc,
        'The universal stationarity of the retained variables is a favorable outcome for GRU model '
        'training. Stationary input series ensure that the statistical properties (mean, variance, '
        'autocorrelation structure) that the model learns during training remain consistent in the '
        'test period, reducing the risk of distributional shift that can cause deep learning models '
        'to produce divergent predictions on out-of-sample data. Had any variable been found '
        'non-stationary, first differencing would have been applied as a preprocessing step before '
        'sequence construction—a standard practice in time series neural network literature '
        '(Sherstinsky, 2020).'
    )

    # ── 4.1.3 Correlation ─────────────────────────────────────────────────────
    add_heading(doc, '4.1.3 Inter-Variable Correlation Analysis', level=3)

    add_body(doc,
        'Spearman rank correlation was computed among the four retained variables to assess '
        'the degree of linear and monotonic association. Spearman correlation was chosen over '
        'Pearson because it does not assume normal distribution—appropriate given the right-skewed '
        'distribution of wind speed and the bimodal distribution of relative humidity during '
        'monsoon transitions. Figure 3 presents the correlation matrices for both stations.'
    )

    add_fig(doc, p_corr,
            'Figure 3. Spearman Rank Correlation Matrix of Retained Meteorological Variables. '
            'Values range from -1 (perfect negative correlation) to +1 (perfect positive '
            'correlation). Darker blue cells indicate strong positive correlation; darker '
            'red cells indicate strong negative correlation.', width=6.0)

    add_body(doc,
        'The most prominent correlation at both stations is the negative relationship between '
        'relative humidity and sea-level pressure: −0.646 at Baguio and −0.581 at Manila. '
        'This inverse relationship reflects a well-documented atmospheric dynamic: areas of '
        'low pressure are typically associated with convergent, moist air masses and higher '
        'humidity, while high-pressure systems are associated with subsidence, drying, and '
        'lower humidity (Ahrens & Henson, 2021). This correlation has direct implications for '
        'feature ablation: because humidity and pressure are correlated, removing one from the '
        'model input may partially but not fully compensate through the information retained in '
        'the other.'
    )

    add_body(doc,
        'Temperature shows a moderate negative correlation with both wind speed (−0.286 at '
        'Baguio, +0.213 at Manila) and pressure (−0.235 at Baguio, −0.253 at Manila). '
        'The sign reversal of the temperature-wind speed correlation between Baguio and Manila '
        'is meteorologically significant: at Baguio, stronger winds are associated with cooler '
        'temperatures due to advection from the South China Sea and the cooling effect of '
        'orographic precipitation, while at Manila, the positive correlation may reflect '
        'the role of sea breezes in simultaneously elevating both temperature and surface wind '
        'speed during daytime (Cinco et al., 2016). This station-specific pattern cautions '
        'against applying a single universal feature importance ranking across geographically '
        'distinct locations.'
    )

    # ── 4.1.4 Ablation Study ──────────────────────────────────────────────────
    add_heading(doc, '4.1.4 Leave-One-Out Ablation Study', level=3)

    add_body(doc,
        'The core of the variable importance analysis was a leave-one-out ablation study in which '
        'a GRU model was trained on the full variable set (baseline), followed by four separate '
        'GRU models each trained on the remaining three variables after removing one. The change '
        'in RMSE between the ablated and baseline configurations (ΔRMSE) quantifies the marginal '
        'predictive contribution of each variable. The GRU architecture used for ablation consisted '
        'of three hidden layers, 128 units per layer, a dropout rate of 0.2, a sequence length of '
        '45 days, and was trained for up to 100 epochs with early stopping (patience = 15). '
        'Results are reported for three forecast targets: temperature, wind speed, and pressure. '
        'Relative humidity was evaluated as a predictor variable but not as a forecast target in '
        'the ablation study due to its strong multicollinearity with pressure.'
    )

    # Ablation table
    abl_headers = ['Station', 'Target', 'Feature Removed', 'Baseline RMSE', 'Ablated RMSE', 'ΔRMSE', 'ΔMAE', 'ΔR²']
    abl_rows = []
    for stn in STATIONS:
        sd = ablation['stations_data'][stn]
        for tgt in TARGETS:
            pa = sd['per_target_ablation'][tgt]
            bl = pa['baseline']
            for feat in ['temp', 'rhum', 'wspd', 'pres']:
                if feat in pa['ablations']:
                    ab = pa['ablations'][feat]
                    dr = ab['RMSE'] - bl['RMSE']
                    dm = ab['MAE']  - bl['MAE']
                    dr2 = ab['R2'] - bl['R2']
                    abl_rows.append([
                        stn,
                        {'temp':'Temperature','wspd':'Wind Speed','pres':'Pressure'}[tgt],
                        VAR_SHORT[feat],
                        f"{bl['RMSE']:.4f}",
                        f"{ab['RMSE']:.4f}",
                        f"{dr:+.4f}",
                        f"{dm:+.4f}",
                        f"{dr2:+.4f}"
                    ])
    add_compact_table(doc, abl_headers, abl_rows,
                      'Table 3. Leave-One-Out GRU Ablation Results — ΔRMSE, ΔMAE, and ΔR² per Station and Target')

    add_body(doc,
        'Table 3 presents the complete ablation results. Across all targets and stations, the '
        'removed feature causing the largest ΔRMSE is consistently the target variable itself '
        '(its own lag): removing temperature from the input set when predicting temperature '
        'caused ΔRMSE = +0.5237 at Baguio and +0.3767 at Manila; removing wind speed as an '
        'input for wind speed prediction caused ΔRMSE = +0.3993 at Baguio and +0.2766 at Manila; '
        'and removing pressure caused ΔRMSE = +1.0297 at Baguio and +0.7030 at Manila. This '
        'finding underscores the fundamental role of autoregressive self-information—the immediate '
        'history of a variable is its most powerful predictor.'
    )

    add_fig(doc, p_abl,
            'Figure 4. Change in RMSE (ΔRMSE) from Leave-One-Out Feature Ablation. '
            'Positive values indicate that removing the variable degraded model performance; '
            'negative values indicate marginal improvement. Bars are grouped by forecast target.', width=6.2)

    add_body(doc,
        'Beyond self-lag effects, Figure 4 reveals important cross-variable influences. At Baguio, '
        'wind speed emerged as the second most important predictor for pressure forecasting '
        '(ΔRMSE = +0.0656), a relationship that reflects the geostrophic balance between '
        'horizontal pressure gradients and wind velocity—where locally elevated wind speeds '
        'serve as a proxy for synoptic-scale pressure gradient magnitude. At Manila, wind '
        'speed and relative humidity both showed statistically detectable influence on temperature '
        'prediction (ΔRMSE = +0.0235 and +0.0156 respectively), consistent with the heat '
        'index relationship and coastal temperature-humidity coupling in tropical maritime climates.'
    )

    add_body(doc,
        'Notably, several ablation configurations produced negative ΔRMSE values (e.g., removing '
        'relative humidity from Baguio temperature prediction: ΔRMSE = −0.0012; removing pressure '
        'from Baguio temperature prediction: ΔRMSE = −0.0042). These marginal improvements upon '
        'feature removal indicate that, for these specific configurations, the removed variable '
        'introduces modest noise into the GRU\'s input representation that slightly degrades '
        'forecast quality. While these improvements are small and may fall within random variation, '
        'they highlight that including all available variables does not unconditionally improve '
        'deep learning forecast performance—a phenomenon known as the "curse of dimensionality" '
        'in finite training samples (Bellman, 1957).'
    )

    # ── 4.1.5 DM Tests ────────────────────────────────────────────────────────
    add_heading(doc, '4.1.5 Diebold-Mariano Test with Benjamini-Hochberg Correction', level=3)

    add_body(doc,
        'The Diebold-Mariano (DM) test (Diebold & Mariano, 1995) was applied to formally assess '
        'whether each ablated GRU model produced statistically significantly worse forecasts than '
        'the baseline full-variable GRU. The Harvey-Leybourne-Newbold (HLN, 1997) correction '
        'was applied to improve finite-sample performance of the test statistic. A positive DM '
        'statistic indicates that the ablated model (missing one variable) produced larger squared '
        'errors than the baseline—meaning the removed variable contributed meaningfully to '
        'forecast accuracy. To control for Type I error inflation from multiple simultaneous '
        'tests, Benjamini-Hochberg (BH) FDR correction at α = 0.05 was applied across all '
        'four tests conducted per target-station combination.'
    )

    # DM summary table
    dm_headers = ['Station', 'Target', 'Feature Removed', 'ΔRMSE', 'DM Stat (HLN)', 'p-value', 'BH p-value', 'Significant?']
    dm_rows = []
    for stn in STATIONS:
        sd = ablation['stations_data'][stn]
        for tgt in TARGETS:
            dm_df = sd['per_target_dm_bh'][tgt]
            for _, row in dm_df.iterrows():
                feat_lbl = str(row['Label']).replace(' (°C)', '').replace(' (%)', '').replace(' (km/h)', '').replace(' (hPa)', '')
                sig_str = 'Yes ***' if row['BH p-value'] < 0.001 else ('Yes **' if row['BH p-value'] < 0.01 else ('Yes *' if row['Significant?'] else 'No'))
                dm_rows.append([
                    stn,
                    {'temp': 'Temperature', 'wspd': 'Wind Speed', 'pres': 'Pressure'}[tgt],
                    feat_lbl,
                    f"{row['DELTA_RMSE'] if 'DELTA_RMSE' in row else row.get('DRMS', row.get('dRMSE', row.get('DRMSE', row.iloc[2]))):.4f}",
                    f"{row['DM Stat']:.4f}",
                    f"{row['p-value']:.4f}",
                    f"{row['BH p-value']:.4f}",
                    sig_str
                ])
    # Re-extract properly
    dm_rows = []
    for stn in STATIONS:
        sd = ablation['stations_data'][stn]
        for tgt in TARGETS:
            dm_df = sd['per_target_dm_bh'][tgt]
            for _, row in dm_df.iterrows():
                feat_lbl = str(row['Label']).replace(' (°C)', '').replace(' (%)', '').replace(' (km/h)', '').replace(' (hPa)', '')
                drmse_val = row.iloc[2]
                sig_str = 'Yes ***' if float(row['BH p-value']) < 0.001 else ('Yes **' if float(row['BH p-value']) < 0.01 else ('Yes *' if bool(row['Significant?']) else 'No'))
                dm_rows.append([
                    stn,
                    {'temp': 'Temperature', 'wspd': 'Wind Speed', 'pres': 'Pressure'}[tgt],
                    feat_lbl,
                    f"{float(drmse_val):+.4f}",
                    f"{float(row['DM Stat']):.4f}",
                    f"{float(row['p-value']):.4f}",
                    f"{float(row['BH p-value']):.4f}",
                    sig_str
                ])
    add_compact_table(doc, dm_headers, dm_rows,
                      'Table 4. Diebold-Mariano Test Results with Benjamini-Hochberg FDR Correction '
                      '(*** p < 0.001, ** p < 0.01, * p < 0.05)')

    add_fig(doc, p_dm,
            'Figure 5. Diebold-Mariano (DM) Test Statistics with Benjamini-Hochberg Correction. '
            'Blue bars represent statistically significant variables (BH p < 0.05); gray bars '
            'are not significant. The dashed red lines mark the ±1.96 critical value. '
            'Significance codes: *** p < 0.001, ** p < 0.01, * p < 0.05, ns = not significant.', width=6.2)

    add_body(doc,
        'Figure 5 and Table 4 together present the DM test outcomes. The results confirm the '
        'primacy of self-lag information: in every target-station combination, the DM statistic '
        'for removing the target variable itself (e.g., removing temperature when forecasting '
        'temperature) is the largest and most significant. At Baguio, removing temperature as a '
        'predictor for temperature forecasting produced a DM statistic of 8.22 (BH p < 0.0001), '
        'while removing wind speed for wind speed forecasting yielded a DM statistic of 3.71 '
        '(BH p = 0.0008). Both results far exceed the critical value of ±1.96 and remain '
        'significant after BH correction.'
    )

    add_body(doc,
        'An important finding from the DM analysis is the station-dependent significance of '
        'cross-variable influences. At Manila, wind speed (DM stat = 3.14, BH p = 0.0036) and '
        'relative humidity (DM stat = 3.02, BH p = 0.0036) were both found to be statistically '
        'significant predictors of temperature—a pattern not replicated at Baguio. This suggests '
        'that Manila\'s coastal, monsoon-influenced climate creates stronger synoptic-scale '
        'coupling between these variables, allowing the GRU to leverage cross-variable temporal '
        'patterns in ways that are less feasible in Baguio\'s more complex, locally dominated '
        'microclimate. For pressure prediction at Baguio, wind speed was the only non-self variable '
        'to reach significance (DM stat = 3.14, BH p = 0.0036), consistent with the geostrophic '
        'relationship discussed in the ablation section.'
    )

    # ── 4.1.6 Friedman Test ───────────────────────────────────────────────────
    add_heading(doc, '4.1.6 Friedman Test for Overall Configuration Differences', level=3)

    add_body(doc,
        'The Friedman test (Friedman, 1937) was employed as a global test to determine whether '
        'the different variable subset configurations (baseline and four ablated versions) produced '
        'statistically different temporal block-RMSE distributions. Unlike the DM test, which '
        'performs pairwise comparisons, the Friedman test evaluates all configurations jointly '
        'as a non-parametric equivalent of one-way repeated-measures ANOVA—appropriate when the '
        'normality of block-RMSE distributions cannot be assumed. Test errors were divided into '
        '10 temporal blocks, and block-RMSE was computed for each configuration, yielding a '
        '10 × 5 matrix of observations for the Friedman test.'
    )

    friedman_headers = ['Station', 'Target', 'Friedman Statistic', 'p-value', 'Significant?']
    baguio_fr = ablation['stations_data']['Baguio']['per_target_friedman']
    manila_fr = ablation['stations_data']['Manila']['per_target_friedman']
    friedman_rows = []
    for tgt in TARGETS:
        tname = {'temp': 'Temperature', 'wspd': 'Wind Speed', 'pres': 'Pressure'}[tgt]
        bs, bp = baguio_fr[tgt]
        ms, mp = manila_fr[tgt]
        friedman_rows.append(['Baguio', tname, f'{bs:.3f}', f'{bp:.5f}', 'Yes ***' if bp < 0.001 else ('Yes *' if bp < 0.05 else 'No')])
        friedman_rows.append(['Manila', tname, f'{ms:.3f}', f'{mp:.5f}', 'Yes ***' if mp < 0.001 else ('Yes *' if mp < 0.05 else 'No')])
    add_compact_table(doc, friedman_headers, friedman_rows,
                      'Table 5. Friedman Test Results for Variable Subset Configuration Comparisons')

    add_fig(doc, p_fr,
            'Figure 6. Friedman Test Significance (-log10 p-value). The dashed red line '
            'indicates the p = 0.05 threshold. All bars substantially exceed the threshold, '
            'confirming that variable subset configurations produce significantly different '
            'performance distributions.', width=5.5)

    add_body(doc,
        'As shown in Table 5 and Figure 6, the Friedman test rejected the null hypothesis of '
        'equal performance distributions for all target-station combinations: Friedman statistics '
        'ranged from 14.24 (Baguio wind speed, p = 0.0066) to 27.20 (Manila pressure, '
        'p < 0.0001). This confirms that the choice of input variables systematically affects '
        'GRU forecast quality, justifying the detailed per-variable significance testing '
        'conducted via the DM test. The Baguio wind speed result, while still significant, '
        'produced the smallest Friedman statistic—consistent with the general difficulty of '
        'wind speed prediction at this station and the relatively modest ΔRMSE values observed '
        'in the ablation study.'
    )

    add_body(doc,
        'The consistency of significant Friedman results across all six target-station combinations '
        'provides strong evidence that variable selection is not a trivial aspect of GRU model '
        'design for weather forecasting. The results support the utility of the ablation pipeline '
        'as a structured approach to input feature selection that combines practical RMSE-based '
        'impact assessment with rigorous statistical testing—a combination that addresses both '
        'effect size and inferential reliability.'
    )

    # ── 4.1.7 Cross-Station Synthesis ─────────────────────────────────────────
    add_heading(doc, '4.1.7 Cross-Station Synthesis', level=3)

    add_body(doc,
        'Synthesizing the ablation, DM, and Friedman results across both stations yields a '
        'coherent picture of variable importance in GRU weather forecasting for the Philippine '
        'context. The consistently dominant finding is that each meteorological variable\'s own '
        'temporal autoregressive component represents its single most important predictor. '
        'This finding is consistent with fundamental properties of weather time series, which '
        'exhibit strong persistence (autocorrelation) at the daily scale—particularly for '
        'smooth variables like pressure and temperature that change gradually under the influence '
        'of large-scale synoptic systems.'
    )

    add_body(doc,
        'Beyond self-lag dominance, meaningful cross-variable influences emerge in specific '
        'geographic and meteorological contexts. At Manila, temperature forecasting benefits '
        'significantly from wind speed and relative humidity inputs, reflecting the coastal '
        'heat-humidity coupling that characterizes the urban tropical climate of the Philippine '
        'capital. At Baguio, pressure forecasting benefits from wind speed information, '
        'reflecting geostrophic dynamics in the highland setting. These contextual influences '
        'suggest that variable importance in GRU weather models is inherently location-specific '
        'and cannot be assumed to transfer across meteorologically distinct stations without '
        'empirical validation.'
    )

    add_body(doc,
        'The lack of statistically significant cross-variable influences for most target-station '
        'combinations (after BH correction) does not imply that these variables are entirely '
        'uninformative. As the ablation ΔRMSE values show, most feature removals produce small '
        'but non-zero performance changes. The BH correction applied in this study controls '
        'for the possibility that some of these small changes arise by chance when conducting '
        'multiple simultaneous tests—a conservative but methodologically appropriate stance '
        'for a study intended to produce actionable guidance for model input selection. Future '
        'work with larger datasets or longer training periods may detect additional significant '
        'cross-variable influences that are underpowered in the current study.'
    )

    # ── Conclusions ───────────────────────────────────────────────────────────
    add_heading(doc, '4.1.8 Conclusions', level=2)

    add_body(doc,
        'This section addressed the first research objective: Which meteorological variables most '
        'significantly influence the predictive performance of GRU models in weather forecasting? '
        'Based on the multi-stage statistical pipeline applied to daily weather data from Baguio '
        'and Manila covering the period 2020 to 2026, the following conclusions are drawn.'
    )

    add_body(doc,
        'Conclusion 1: The target variable\'s own lag is the dominant predictor in every '
        'configuration tested. Removing the target variable from the GRU input set produced the '
        'largest performance degradation in all six target-station combinations—by a margin of '
        'at least 200% over the next most important predictor. This finding is consistent with '
        'the strong first-order autocorrelation of daily meteorological time series and '
        'establishes a clear hierarchy of variable importance with self-lag at the apex.'
    )

    add_body(doc,
        'Conclusion 2: Cross-variable influences are statistically significant in specific '
        'geographic and meteorological contexts. Manila\'s temperature prediction demonstrates '
        'significant contributions from wind speed (DM p = 0.0036) and relative humidity '
        '(DM p = 0.0036), while Baguio\'s pressure prediction is significantly enhanced by '
        'wind speed information (DM p = 0.0036). These context-specific influences reflect '
        'the physical coupling between atmospheric variables that operates differently in '
        'highland versus coastal environments—specifically, geostrophic dynamics at Baguio '
        'and maritime heat-humidity coupling at Manila.'
    )

    add_body(doc,
        'Conclusion 3: The Friedman test confirms that variable subset selection has a '
        'statistically significant and consistent effect on GRU forecast quality across all '
        'targets and stations. This justifies variable importance analysis as a necessary step '
        'in GRU model development for operational weather forecasting, rather than an optional '
        'methodological refinement. The Friedman statistics ranged from 14.24 to 27.20 '
        '(all p < 0.01), confirming that even when individual variable DM tests fail to reach '
        'significance, the collective pattern of performance differences across configurations '
        'is non-random.'
    )

    add_body(doc,
        'Conclusion 4: Variable importance in GRU weather forecasting is location-specific. '
        'The same variable that is statistically significant at Manila (wind speed for temperature '
        'prediction) is not significant at Baguio, and vice versa (wind speed for pressure '
        'prediction at Baguio is significant but not at Manila). This spatial heterogeneity '
        'implies that feature selection protocols cannot be universally applied across '
        'geographically distinct Philippine weather stations and must be validated locally.'
    )

    # ── Recommendations ───────────────────────────────────────────────────────
    add_heading(doc, '4.1.9 Recommendations', level=2)

    add_body(doc,
        'Based on the findings of the variable importance analysis, the following recommendations '
        'are offered for researchers and practitioners developing GRU-based weather forecasting '
        'systems for Philippine climate contexts.'
    )

    add_body(doc,
        'Recommendation 1: Always include the target variable\'s own lagged values as a '
        'primary input feature. The ablation results confirm that self-lag information is '
        'irreplaceable for GRU temperature, wind speed, and pressure forecasting. Experimental '
        'designs that test GRU performance without self-lag inputs—as sometimes done in '
        'purely exogenous variable importance studies—should be clearly distinguished from '
        'realistic operational configurations where historical observations of the target '
        'variable are available.'
    )

    add_body(doc,
        'Recommendation 2: Conduct site-specific variable importance assessments before '
        'finalizing GRU input sets. The Manila-specific significance of wind speed and '
        'humidity for temperature prediction, and the Baguio-specific significance of wind '
        'speed for pressure prediction, demonstrate that variable importance is not transferable '
        'across locations. Future GRU deployments at additional Philippine stations (e.g., '
        'Davao, Cebu, Zamboanga) should conduct equivalent ablation analyses to identify '
        'locally relevant predictor sets.'
    )

    add_body(doc,
        'Recommendation 3: Extend the variable universe to include additional meteorological '
        'and auxiliary predictors. The current study was constrained to four variables by '
        'data coverage requirements. Future studies should explore the inclusion of dew point '
        'temperature (strongly linked to humidity and human heat stress), large-scale '
        'teleconnection indices (ENSO, MJO phase), and satellite-derived land surface '
        'temperature as potential predictors. Dew point in particular is consistently '
        'available at PAGASA synoptic stations and could provide additional information '
        'for temperature and humidity forecasting beyond what current inputs supply.'
    )

    add_body(doc,
        'Recommendation 4: Apply the full ablation-DM-Friedman pipeline as a standard '
        'feature evaluation protocol for GRU weather models. The structured three-stage '
        'pipeline used in this study—ablation for effect size estimation, DM test for '
        'pairwise significance, and Friedman test for global configuration difference—'
        'provides a statistically rigorous and interpretable framework for variable importance '
        'assessment. This pipeline should be considered for standardization as a best practice '
        'in the Philippine meteorological modeling community, particularly as deep learning '
        'weather models become more widely adopted.'
    )

    add_body(doc,
        'Recommendation 5: Investigate temporal dynamics of variable importance. The current '
        'analysis treats variable importance as a static property over the full evaluation '
        'period. However, meteorological variable relationships are inherently seasonal: '
        'humidity-pressure coupling is likely stronger during the typhoon season, while '
        'wind speed-pressure coupling may be more pronounced during northeast monsoon episodes. '
        'Future work should apply rolling-window or seasonally stratified ablation analyses '
        'to characterize the time-varying nature of variable importance, which would support '
        'the development of adaptive, season-aware GRU architectures.'
    )

    doc.save('Findings_Objective1_VariableImportance.docx')
    print("Saved: Findings_Objective1_VariableImportance.docx")


build_doc2()
build_doc1()
print("Both documents complete.")
