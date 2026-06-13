"""
generate_documentation.py
Generates system_documentation.docx — comprehensive technical documentation
covering data pipeline, model architectures, metrics (with worked examples),
variable importance analysis, and the web application.

Run:  python generate_documentation.py
"""

import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Palette ────────────────────────────────────────────────────────────────────
def _hex(h: str) -> RGBColor:
    b = bytes.fromhex(h)
    return RGBColor(b[0], b[1], b[2])

C_NAVY  = "1f3864"
C_BLUE  = "2471a3"
C_TEAL  = "145a32"
C_PURP  = "4a235a"
C_DARK  = "2c3e50"
C_LIGHT = "eaf0fb"
C_WHITE = "ffffff"
C_STEEL = "adb5bd"
C_MGRAY = "f4f6f7"
C_WARN  = "fef9e7"
C_CODE  = "f0f3f4"


# ── Low-level helpers (reused from generate_docs.py) ──────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def no_borders(cell):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tcBdr.append(el)
    tcPr.append(tcBdr)


def cell_margin(cell, top=60, bottom=60, left=120, right=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _no_tbl_borders(tbl):
    tbl_el = tbl._tbl
    tblPr  = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBdr.append(el)
    tblPr.append(tblBdr)


# ── Document-level helpers ─────────────────────────────────────────────────────

def sp(doc, before=4, after=4):
    """Empty spacer paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    return p


def h1(doc, text: str):
    """Top-level section heading — navy, 15 pt bold, with coloured left bar."""
    tbl = doc.add_table(rows=1, cols=2)
    _no_tbl_borders(tbl)
    tbl.columns[0].width = Cm(0.25)
    tbl.columns[1].width = Cm(16.25)
    bar  = tbl.rows[0].cells[0]
    main = tbl.rows[0].cells[1]
    set_cell_bg(bar, C_BLUE); no_borders(bar)
    no_borders(main)
    cell_margin(main, top=50, bottom=50, left=160, right=40)
    p = main.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(15)
    r.font.color.rgb = _hex(C_NAVY)
    sp(doc, 4, 2)


def h2(doc, text: str):
    """Sub-section heading — dark blue, 12 pt bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = _hex(C_BLUE)


def h3(doc, text: str):
    """Sub-sub-section heading — dark, 11 pt bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = _hex(C_DARK)


def body(doc, text: str, size=10.5, italic=False, color=C_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(size); r.italic = italic
    r.font.color.rgb = _hex(color)
    return p


def bullet(doc, text: str, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(size)


def formula(doc, text: str):
    """Centred formula in a light-blue box."""
    tbl = doc.add_table(rows=1, cols=1)
    _no_tbl_borders(tbl)
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, C_LIGHT); no_borders(c)
    cell_margin(c, top=80, bottom=80, left=200, right=200)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = _hex(C_NAVY)
    sp(doc, 2, 2)


def code_block(doc, text: str):
    """Monospace code block in a gray box."""
    tbl = doc.add_table(rows=1, cols=1)
    _no_tbl_borders(tbl)
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, C_CODE); no_borders(c)
    cell_margin(c, top=80, bottom=80, left=160, right=160)
    for i, line in enumerate(text.strip().split("\n")):
        p = c.paragraphs[0] if i == 0 else c.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(line)
        r.font.size = Pt(9)
        r.font.name = "Courier New"
        r.font.color.rgb = _hex("2c3e50")
    sp(doc, 2, 4)


def note_box(doc, text: str, color=C_WARN, text_color=C_DARK):
    """Highlighted note / callout box."""
    tbl = doc.add_table(rows=1, cols=1)
    _no_tbl_borders(tbl)
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, color); no_borders(c)
    cell_margin(c, top=70, bottom=70, left=160, right=160)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(10); r.italic = True
    r.font.color.rgb = _hex(text_color)
    sp(doc, 2, 4)


def data_table(doc, headers: list, rows: list, col_widths=None,
               header_color=C_NAVY):
    """Formatted table with coloured header row."""
    n = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n)
    tbl.style = "Table Grid"

    if col_widths:
        for i, w in enumerate(col_widths):
            tbl.columns[i].width = Cm(w)

    # Header
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]
        set_cell_bg(c, header_color)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = _hex(C_WHITE)

    # Data rows
    for i, row in enumerate(rows):
        bg = C_LIGHT if i % 2 == 0 else C_WHITE
        for j, val in enumerate(row):
            c = tbl.rows[i + 1].cells[j]
            set_cell_bg(c, bg)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)

    sp(doc, 4, 6)
    return tbl


def section_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), C_STEEL)
    pBdr.append(bot)
    pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════════
# BUILD
# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    for sec in doc.sections:
        sec.page_width    = Cm(21.59)
        sec.page_height   = Cm(27.94)
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.0)

    TODAY = datetime.date.today().strftime("%B %d, %Y")

    # ══════════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════════════════════
    sp(doc, 30, 0)
    tbl = doc.add_table(rows=1, cols=1)
    _no_tbl_borders(tbl)
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, C_NAVY); no_borders(c)
    cell_margin(c, top=200, bottom=200, left=300, right=300)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run("SYSTEM DOCUMENTATION")
    r.bold = True; r.font.size = Pt(22)
    r.font.color.rgb = _hex(C_WHITE)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(6)
    r2 = p2.add_run(
        "Evaluating the Performance of Gated Recurrent Units\n"
        "for Multi-Parameter Weather Forecasting"
    )
    r2.font.size = Pt(13); r2.font.color.rgb = _hex("adc8e0")
    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(0)
    r3 = p3.add_run(
        "Castro, E. M.  ·  Millan, D. J. S.  ·  Mondoñedo, J. E.\n"
        "University of the Cordilleras  ·  College of CITCS  ·  B.S. Computer Science\n"
        f"{TODAY}"
    )
    r3.font.size = Pt(10); r3.font.color.rgb = _hex(C_STEEL)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. SYSTEM OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "1. System Overview")
    body(doc,
         "This system is a web-based benchmarking platform that fetches real daily weather "
         "data from the Meteostat API and evaluates five machine learning and statistical "
         "models for multi-variable weather forecasting across two Philippine stations. "
         "The primary research question is: does the Gated Recurrent Unit (GRU) outperform "
         "LSTM, SimpleRNN, Linear Regression, and ARIMA for multi-parameter weather "
         "forecasting on Philippine station-level data?")

    h2(doc, "1.1  End-to-End Pipeline")
    data_table(doc,
        ["Stage", "Module", "Description"],
        [
            ["1. Data Fetch",    "data_fetch.py",      "Pull daily records from Meteostat API (2020-01-01 → today)"],
            ["2. Preprocess",    "preprocessing.py",   "Impute missing values, IQR outlier removal, MinMax/Log scaling, 80/20 split, sliding-window sequences"],
            ["3. Train Models",  "models.py",          "GRU, LSTM, SimpleRNN (PyTorch), Linear Regression (scikit-learn), ARIMA (pmdarima)"],
            ["4. Evaluate",      "evaluation.py",      "Compute Accuracy, MSE, MAE, R², t-stat / p-value per model per variable"],
            ["5. Visualise",     "model_evaluation.py","Interactive Plotly charts, metrics tables, bar charts"],
            ["6. Forecast",      "model_evaluation.py","Rolling 7-day forecast from last 45-day window"],
            ["7. Var. Importance","feature_analysis.py","Ablation, Diebold-Mariano, Friedman, Benjamini-Hochberg"],
        ],
        col_widths=[3.5, 4.0, 9.0]
    )

    h2(doc, "1.2  Station Reference")
    data_table(doc,
        ["Station", "Meteostat ID", "Location", "Elevation", "Climate Type"],
        [
            ["Baguio City", "98328", "Luzon Cordillera, Philippines", "~1,500 m asl", "Cool, wet highland"],
            ["Manila",      "98425", "National Capital Region, Philippines", "~15 m asl",    "Tropical lowland"],
        ],
        col_widths=[3.0, 3.0, 5.0, 2.5, 3.0]
    )

    h2(doc, "1.3  Technology Stack")
    data_table(doc,
        ["Component", "Library / Tool", "Version / Notes"],
        [
            ["Language",         "Python",             "3.10+"],
            ["Deep Learning",    "PyTorch",            "2.x  —  CUDA → MPS → CPU auto-detected"],
            ["ARIMA",            "pmdarima",           "auto_arima, AIC criterion"],
            ["Classical ML",     "scikit-learn",       "LinearRegression, MinMaxScaler"],
            ["Data API",         "Meteostat",          "Open weather API, daily resolution"],
            ["Data Wrangling",   "pandas / NumPy",     "—"],
            ["Web UI",           "Streamlit",          "1.x, multi-page via st.navigation()"],
            ["Interactive Charts","Plotly",            "6.x, st.plotly_chart()"],
            ["Static Charts",    "Matplotlib",         "Variable Importance page"],
            ["Document Gen.",    "python-docx / pptx", "Thesis documents and slides"],
        ],
        col_widths=[3.5, 3.5, 9.5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. DATA ACQUISITION
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "2. Data Acquisition  (data_fetch.py)")
    body(doc,
         "Weather data is sourced from the Meteostat Open Weather API, which provides "
         "historical daily meteorological records for weather stations worldwide. "
         "Data is fetched dynamically at runtime from January 1, 2020 through today's "
         "date, yielding approximately 2,190 daily records per station.")

    h2(doc, "2.1  Available Meteostat Variables")
    data_table(doc,
        ["Key", "Full Name", "Unit", "Notes"],
        [
            ["temp", "Average Temperature",  "°C",    "Primary forecast target"],
            ["dwpt", "Dew Point",            "°C",    "—"],
            ["rhum", "Relative Humidity",    "%",     "—"],
            ["prcp", "Precipitation",        "mm",    "Primary forecast target — right-skewed"],
            ["snow", "Snow Depth",           "mm",    "Often sparse / missing in PH"],
            ["wdir", "Wind Direction",       "°",     "—"],
            ["wspd", "Wind Speed",           "km/h",  "Primary forecast target"],
            ["wpgt", "Peak Wind Gust",       "km/h",  "—"],
            ["pres", "Sea-level Pressure",   "hPa",   "Primary forecast target"],
            ["tsun", "Sunshine Duration",    "min",   "Often missing in PH stations"],
        ],
        col_widths=[2.0, 4.5, 2.0, 8.0]
    )

    h2(doc, "2.2  Fetch Code")
    code_block(doc, """\
from data_fetch import fetch_all_data
from datetime import datetime

data = fetch_all_data(
    start   = datetime(2020, 1, 1),
    end     = datetime.now(),
    columns = ["prcp", "temp", "wspd", "pres"],
)
# Returns: {"Baguio": pd.DataFrame, "Manila": pd.DataFrame}
# Index: DatetimeIndex  |  Columns: selected variables""")

    note_box(doc,
             "Columns absent from a station are added as all-NaN so downstream code always "
             "receives the full schema regardless of data availability per station.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. DATA PREPROCESSING
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "3. Data Preprocessing  (preprocessing.py)")

    # ── 3.1 Missing Values ────────────────────────────────────────────────────
    h2(doc, "3.1  Missing Value Imputation")
    body(doc,
         "Missing values are filled in three sequential passes to handle both isolated "
         "and consecutive gaps:")
    bullet(doc, "Step 1 — Forward-fill (ffill): propagates last known value forward.")
    bullet(doc, "Step 2 — Backward-fill (bfill): fills any remaining gaps at the start.")
    bullet(doc, "Step 3 — Linear interpolation: resolves any residual NaNs.")

    h3(doc, "Worked Example")
    data_table(doc,
        ["Date", "Raw temp (°C)", "After ffill", "After bfill", "After interpolate"],
        [
            ["2023-06-01", "22.1", "22.1", "22.1", "22.1"],
            ["2023-06-02", "NaN",  "22.1", "22.1", "22.1"],
            ["2023-06-03", "NaN",  "22.1", "22.1", "22.1"],
            ["2023-06-04", "23.5", "23.5", "23.5", "23.5"],
        ],
        col_widths=[3.0, 3.0, 3.0, 3.0, 3.5]
    )
    note_box(doc,
             "Linear interpolation would have yielded the same values here (22.1, 22.5, 23.0) "
             "for a two-day gap — the three passes together handle all edge cases robustly.")

    # ── 3.2 IQR Outlier Removal ───────────────────────────────────────────────
    h2(doc, "3.2  IQR Outlier Removal")
    body(doc,
         "Values outside the interquartile fence are replaced with NaN and then "
         "re-imputed. This preserves the real distribution while clipping measurement errors.")

    formula(doc,
            "Lower fence  =  Q1  −  1.5 × IQR\n"
            "Upper fence  =  Q3  +  1.5 × IQR\n"
            "where  IQR  =  Q3  −  Q1")

    h3(doc, "Worked Example — Temperature (Baguio City)")
    body(doc, "Given a sample of daily temperatures:", size=10.5)
    data_table(doc,
        ["Statistic", "Value"],
        [
            ["Q1  (25th percentile)", "17.2 °C"],
            ["Q3  (75th percentile)", "23.8 °C"],
            ["IQR  =  Q3 − Q1",      "6.6 °C"],
            ["Lower fence  =  17.2 − 1.5 × 6.6", "7.3 °C"],
            ["Upper fence  =  23.8 + 1.5 × 6.6", "33.7 °C"],
            ["Values outside [7.3, 33.7]", "→ replaced with NaN, then re-imputed"],
        ],
        col_widths=[8.0, 8.5]
    )

    note_box(doc,
             "Precipitation (prcp) is intentionally excluded from IQR removal. "
             "Typhoon-driven rainfall can legitimately exceed 200 mm/day. "
             "Fencing would cap extreme values at ~12–15 mm, causing models to "
             "systematically under-predict rainfall during typhoons.")

    # ── 3.3 LogMinMaxScaler ───────────────────────────────────────────────────
    h2(doc, "3.3  Feature Scaling — LogMinMaxScaler (Precipitation)")
    body(doc,
         "Standard MinMax scaling collapses precipitation: on a typical dataset, "
         "10 mm would map to 0.03 while 0 mm maps to 0.00, making rainfall nearly "
         "indistinguishable in scaled space. A log1p transform is applied first:")

    formula(doc,
            "scaled  =  MinMax( log1p( clip(x, 0) ) )\n\n"
            "Inverse:  original  =  expm1( MinMax⁻¹(scaled) )")

    h3(doc, "Sample Transformation Table")
    data_table(doc,
        ["Raw prcp (mm)", "log1p(x)", "Scaled [0, 1]  (example range 0–200 mm)"],
        [
            ["0",   "0.000", "0.000"],
            ["1",   "0.693", "0.128"],
            ["5",   "1.792", "0.331"],
            ["10",  "2.398", "0.443"],
            ["50",  "3.932", "0.727"],
            ["100", "4.615", "0.853"],
            ["200", "5.303", "0.980"],
        ],
        col_widths=[4.0, 4.0, 8.5]
    )

    body(doc,
         "All other variables (temp, wspd, pres) use standard MinMaxScaler "
         "with no log transform. All scalers are fitted on the training set only "
         "and saved for inverse-transforming predictions back to original units.")

    code_block(doc, """\
class LogMinMaxScaler:
    def fit_transform(self, X):
        return self._sc.fit_transform(np.log1p(np.clip(X, 0, None)))

    def inverse_transform(self, X):
        return np.expm1(self._sc.inverse_transform(X))""")

    # ── 3.4 Train/Test Split ──────────────────────────────────────────────────
    h2(doc, "3.4  Temporal Train / Test Split")
    body(doc,
         "Data is split 80% / 20% preserving temporal order — no shuffling. "
         "This ensures the model is always evaluated on future (unseen) data.")

    h3(doc, "Example — Baguio City (2020-01-01 to 2025-06-12)")
    data_table(doc,
        ["Set", "Records", "Approx. Date Range", "Purpose"],
        [
            ["Training", "~1,753 rows", "2020-01-01 → 2024-05-20", "Model fitting"],
            ["Test",     "~438 rows",  "2024-05-21 → 2025-06-12", "Evaluation (unseen)"],
        ],
        col_widths=[2.5, 3.0, 5.0, 5.0]
    )

    # ── 3.5 Sliding Window ────────────────────────────────────────────────────
    h2(doc, "3.5  Sliding-Window Sequences")
    body(doc,
         "RNN-based models require fixed-length input sequences. A sliding window "
         "of seq_len=45 days is applied to produce (X, y) training pairs.")

    formula(doc,
            "X[i]  =  scaled_data[i : i+45]    shape: (45, 4)\n"
            "y[i]  =  scaled_data[i+45]         shape: (4,)\n"
            "N     =  total_rows − seq_len")

    h3(doc, "Visual Diagram")
    data_table(doc,
        ["Window", "Input  X  (rows)", "Target  y  (row)"],
        [
            ["Window 0",  "rows  0 – 44",  "row 45"],
            ["Window 1",  "rows  1 – 45",  "row 46"],
            ["Window 2",  "rows  2 – 46",  "row 47"],
            ["...",       "...",            "..."],
            ["Window N",  "rows N – N+44", "row N+45"],
        ],
        col_widths=[3.5, 6.0, 7.0]
    )

    body(doc,
         "Final tensor shapes: X_train = (N_train, 45, 4),  y_train = (N_train, 4). "
         "Linear Regression flattens X to (N, 180) before fitting. "
         "ARIMA operates on univariate series and does not use sliding windows.")

    code_block(doc, """\
def create_sequences(data, seq_len=45):
    X, y = [], []
    for i in range(len(data) - seq_len):
        X.append(data[i : i + seq_len])   # shape (45, 4)
        y.append(data[i + seq_len])        # shape (4,)
    return np.array(X, float32), np.array(y, float32)""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. MODEL ARCHITECTURES
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "4. Model Architectures  (models.py)")

    # ── 4.1 GRU ──────────────────────────────────────────────────────────────
    h2(doc, "4.1  Gated Recurrent Unit (GRU)")
    body(doc,
         "GRU reduces the vanishing gradient problem through two learned gates that "
         "control how much of the previous hidden state to retain. Unlike LSTM it has "
         "no separate cell state, making it computationally lighter.")

    h3(doc, "Gate Equations (single layer, single direction)")
    formula(doc,
            "Update gate:   zₜ = σ( W₄ xₜ + U₄ hₜ₋₁ + b₄ )\n"
            "Reset gate:    rₜ = σ( Wᵣ xₜ + Uᵣ hₜ₋₁ + bᵣ )\n"
            "Candidate:     h̃ₜ = tanh( Wₕ xₜ + rₜ ⊙ Uₕ hₜ₋₁ + bₕ )\n"
            "Output:        hₜ = (1 − zₜ) ⊙ hₜ₋₁ + zₜ ⊙ h̃ₜ")

    body(doc,
         "The update gate zₜ decides how much of the previous state to keep. "
         "When zₜ ≈ 1 the unit mostly copies the old state (long memory). "
         "The reset gate rₜ controls how much of the past state influences the candidate. "
         "When rₜ ≈ 0 the candidate is computed as if starting fresh.")

    h3(doc, "Architecture in This System")
    data_table(doc,
        ["Layer", "Type", "Output Shape", "Notes"],
        [
            ["Input",       "Sequence tensor",              "(batch, 45, 4)",   "45-day window, 4 variables"],
            ["BiGRU Layer 1","Bidirectional GRU",           "(batch, 45, 512)", "hidden=256 × 2 directions"],
            ["Dropout",     "Dropout(p=0.3)",               "(batch, 45, 512)", "Applied between layers"],
            ["BiGRU Layer 2","Bidirectional GRU",           "(batch, 512)",     "Last time-step extracted"],
            ["Dense 128",   "Linear(512, 128) + ReLU",      "(batch, 128)",     "—"],
            ["Dropout",     "Dropout(p=0.3)",               "(batch, 128)",     "—"],
            ["Output",      "Linear(128, 4)",               "(batch, 4)",       "prcp, temp, wspd, pres"],
        ],
        col_widths=[3.0, 4.5, 3.5, 5.5]
    )

    # ── 4.2 LSTM ──────────────────────────────────────────────────────────────
    h2(doc, "4.2  Long Short-Term Memory (LSTM)")
    body(doc,
         "LSTM uses a dedicated cell state Cₜ alongside the hidden state hₜ, "
         "controlled by three gates. The explicit memory cell allows LSTM to retain "
         "information over longer sequences than a plain RNN.")

    h3(doc, "Gate Equations")
    formula(doc,
            "Forget gate:  fₜ = σ( Wₑ [hₜ₋₁, xₜ] + bₑ )\n"
            "Input gate:   iₜ = σ( Wᵢ [hₜ₋₁, xₜ] + bᵢ )\n"
            "Cell update:  C̃ₜ = tanh( Wᶜ [hₜ₋₁, xₜ] + bᶜ )\n"
            "Cell state:   Cₜ = fₜ ⊙ Cₜ₋₁ + iₜ ⊙ C̃ₜ\n"
            "Output gate:  oₜ = σ( Wₒ [hₜ₋₁, xₜ] + bₒ )\n"
            "Hidden:       hₜ = oₜ ⊙ tanh(Cₜ)")

    note_box(doc,
             "GRU vs LSTM: GRU merges the forget and input gates into a single update gate "
             "and has no cell state — 2 gates total vs LSTM's 3 gates + cell state. "
             "GRU is faster to train; LSTM may model very long-range dependencies better. "
             "Same bidirectional 2-layer architecture is used for both in this system.")

    # ── 4.3 SimpleRNN ─────────────────────────────────────────────────────────
    h2(doc, "4.3  Simple Recurrent Neural Network (SimpleRNN)")
    body(doc,
         "The vanilla RNN updates the hidden state with a single tanh activation. "
         "It serves as a baseline for the gated architectures.")

    formula(doc, "hₜ = tanh( Wₕ xₜ + Uₕ hₜ₋₁ + b )")

    body(doc,
         "SimpleRNN is run unidirectional only. Bidirectional plain RNNs are "
         "numerically unstable over long sequences because gradients must backpropagate "
         "through two parallel chains without gating. Hidden=256, 2 layers, dropout=0.3.")

    note_box(doc,
             "Vanishing gradient: in a plain RNN, gradients are multiplied by Wₕ at every "
             "time step. For sequences of length 45, a weight matrix with spectral radius < 1 "
             "causes gradients to decay exponentially, making it hard to learn dependencies "
             "beyond ~10 steps. GRU and LSTM gates act as identity paths that preserve gradients.")

    # ── 4.4 Linear Regression ─────────────────────────────────────────────────
    h2(doc, "4.4  Linear Regression")
    body(doc,
         "The 45-day × 4-variable input window is flattened to a 180-dimensional "
         "vector before fitting a standard Ordinary Least Squares regression.")

    formula(doc,
            "Flatten:  Xₕₑₐₜ = reshape(X, (N, 180))\n"
            "Fit:      β = (XᵀX)⁻¹ Xᵀ y\n"
            "Predict:  ŷ = Xᵀ β")

    body(doc,
         "Linear Regression assumes linear relationships between inputs and outputs. "
         "It acts as a classical statistical baseline: if deep learning models do not "
         "significantly outperform it, the non-linear complexity may not be justified.")

    code_block(doc, """\
def train_linear_regression(X_train, y_train):
    lr = LinearRegression()
    lr.fit(X_train.reshape(len(X_train), -1), y_train)  # flatten (N, 45, 4) -> (N, 180)
    return lr""")

    # ── 4.5 ARIMA ─────────────────────────────────────────────────────────────
    h2(doc, "4.5  ARIMA — Auto-Regressive Integrated Moving Average")
    body(doc,
         "ARIMA(p, d, q) models are fitted independently per variable using pmdarima's "
         "auto_arima with AIC-based stepwise order selection.")

    formula(doc,
            "AR(p):  yₜ = φ₁yₜ₋₁ + φ₂yₜ₋₂ + ... + φ₟yₜ₋₟ + εₜ\n"
            "I(d):   d-th difference removes trend (d=0: no diff, d=1: first diff)\n"
            "MA(q):  εₜ + θ₁εₜ₋₁ + ... + θₓεₜ₋ₓ  (moving average of errors)")

    h3(doc, "Rolling 1-Step-Ahead Evaluation Strategy")
    body(doc,
         "Bulk long-horizon forecasting (predict N steps at once) causes ARIMA to "
         "diverge toward its long-run mean or into negative values for integrated "
         "series (d ≥ 1). Instead, this system uses rolling 1-step-ahead evaluation:")

    data_table(doc,
        ["Step", "Action"],
        [
            ["1", "Deep-copy the fitted ARIMA model (preserves original for 7-day forecast)"],
            ["2", "Update the copy with the first seq_len=45 test observations as context"],
            ["3", "Predict 1 step ahead"],
            ["4", "Feed the true test value back (teacher forcing equivalent)"],
            ["5", "Repeat steps 3–4 for all n_out = len(test) − 45 steps"],
        ],
        col_widths=[1.5, 15.0]
    )

    code_block(doc, """\
def predict_arima(arima_models, test_df_len, seq_len, columns, test_series):
    for j, col in enumerate(columns):
        m = copy.deepcopy(arima_models[col])        # isolate: don't mutate original
        m.update(test_series[:seq_len, j].tolist(), refit=False)
        for i in range(n_out):
            fc = float(np.clip(m.predict(n_periods=1)[0], 0.0, 1.0))
            preds.append(fc)
            m.update([float(test_series[seq_len + i, j])], refit=False)""")

    note_box(doc,
             "refit=False uses Kalman-filter state updates (incremental) instead of full "
             "MLE re-estimation, making rolling evaluation fast. The deep copy ensures "
             "the original model stored for 7-day forecasting is never mutated.")

    # ── 4.6 Training Loop ─────────────────────────────────────────────────────
    h2(doc, "4.6  Deep Learning Training Loop")
    data_table(doc,
        ["Component", "Setting", "Purpose"],
        [
            ["Optimizer",         "Adam  (lr = 1e-3)",            "Adaptive gradient descent"],
            ["LR Scheduler",      "ReduceLROnPlateau  (factor=0.5, patience=5)", "Decay LR when val loss plateaus"],
            ["Loss Function",     "MSE",                          "Mean Squared Error on scaled predictions"],
            ["Gradient Clipping", "max_norm = 1.0",               "Prevents exploding gradients on long sequences"],
            ["Early Stopping",    "patience = 30",                "Halt training when val MSE stops improving"],
            ["Best Weights",      "Restored on stop",             "Returns the checkpoint with lowest val MSE"],
        ],
        col_widths=[3.5, 5.5, 7.5]
    )

    code_block(doc, """\
optimizer = torch.optim.Adam(model.parameters(), lr=1e-3)
scheduler = ReduceLROnPlateau(optimizer, mode='min', factor=0.5, patience=5)

for epoch in range(1, epochs + 1):
    # --- Training pass ---
    for X_batch, y_batch in loader:
        optimizer.zero_grad()
        loss = criterion(model(X_batch), y_batch)
        loss.backward()
        nn.utils.clip_grad_norm_(model.parameters(), max_norm=1.0)
        optimizer.step()

    # --- Validation pass ---
    val_loss = criterion(model(X_val), y_val).item()
    scheduler.step(val_loss)

    if val_loss < best_val_loss:
        best_val_loss = val_loss
        best_state    = copy of model weights
        patience_ctr  = 0
    else:
        patience_ctr += 1
        if patience_ctr >= patience:
            break  # early stop

model.load_state_dict(best_state)  # restore best checkpoint""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. EVALUATION METRICS
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "5. Evaluation Metrics  (evaluation.py)")
    body(doc,
         "Six metrics are computed per (model, variable) pair on the test set "
         "after inverse-transforming predictions back to original physical units.")

    # ── 5.1 Accuracy ──────────────────────────────────────────────────────────
    h2(doc, "5.1  Accuracy (%)")
    body(doc,
         "Accuracy is derived from MAPE (Mean Absolute Percentage Error) and is the "
         "primary interpretability metric. Only days where |y_true| ≥ 1.0 are included "
         "to avoid blow-up from trace precipitation or near-zero values.")

    formula(doc,
            "MAPE  =  (1/N) × Σ |yᵢ − ŷᵢ| / |yᵢ|  × 100%    (where |yᵢ| ≥ 1.0)\n\n"
            "Accuracy  =  max(0,  100 − MAPE)")

    h3(doc, "Worked Example — Temperature Prediction")
    data_table(doc,
        ["Day", "Actual y (°C)", "Predicted ŷ (°C)", "|y − ŷ| / |y|  × 100%"],
        [
            ["1",  "22.0", "21.3", "3.18%"],
            ["2",  "25.5", "24.1", "5.49%"],
            ["3",  "18.8", "20.0", "6.38%"],
            ["4",  "27.2", "26.9", "1.10%"],
            ["5",  "23.1", "22.8", "1.30%"],
            ["",   "",     "MAPE  =  (3.18+5.49+6.38+1.10+1.30)/5", "= 3.49%"],
            ["",   "",     "Accuracy  =  100 − 3.49",               "= 96.51%"],
        ],
        col_widths=[1.5, 3.5, 3.5, 8.0]
    )

    # ── 5.2 MSE ───────────────────────────────────────────────────────────────
    h2(doc, "5.2  Mean Squared Error (MSE)")
    formula(doc, "MSE  =  (1/N)  ×  Σ (yᵢ − ŷᵢ)²")
    body(doc,
         "MSE penalises large errors heavily due to squaring. It is particularly "
         "sensitive to outlier predictions, which is relevant for extreme precipitation events.")

    h3(doc, "Worked Example (same 5 days)")
    data_table(doc,
        ["Day", "Error  (y − ŷ)", "Squared Error  (y − ŷ)²"],
        [
            ["1", " 0.70", " 0.490"],
            ["2", " 1.40", " 1.960"],
            ["3", "-1.20", " 1.440"],
            ["4", " 0.30", " 0.090"],
            ["5", " 0.30", " 0.090"],
            ["", "Sum", "4.070"],
            ["", "MSE  =  4.070 / 5", "= 0.814 °C²"],
        ],
        col_widths=[1.5, 5.0, 10.0]
    )

    # ── 5.3 MAE ───────────────────────────────────────────────────────────────
    h2(doc, "5.3  Mean Absolute Error (MAE)")
    formula(doc, "MAE  =  (1/N)  ×  Σ |yᵢ − ŷᵢ|")
    body(doc,
         "MAE is robust to outliers and directly interpretable in the original unit "
         "(mm for precipitation, °C for temperature, etc.).")

    h3(doc, "Worked Example (same 5 days)")
    body(doc, "MAE  =  (0.70 + 1.40 + 1.20 + 0.30 + 0.30) / 5  =  3.90 / 5  =  0.78 °C", italic=True)

    # ── 5.4 R² ────────────────────────────────────────────────────────────────
    h2(doc, "5.4  Coefficient of Determination (R²)")
    formula(doc,
            "R²  =  1  −  SSᵣₑₛ / SSₜₒₜ\n\n"
            "SSᵣₑₛ  =  Σ (yᵢ − ŷᵢ)²     (residual sum of squares)\n"
            "SSₜₒₜ  =  Σ (yᵢ − ȳ)²      (total sum of squares,  ȳ = mean of y)")

    h3(doc, "Worked Example (same 5 days)")
    body(doc, "y = [22.0, 25.5, 18.8, 27.2, 23.1]  →  ȳ = 23.32 °C", italic=True)
    data_table(doc,
        ["Day", "(y − ŷ)²", "(y − ȳ)²"],
        [
            ["1", "0.490", "1.742"],
            ["2", "1.960", "4.756"],
            ["3", "1.440", "20.684"],
            ["4", "0.090", "14.824"],
            ["5", "0.090", "0.049"],
            ["Sum", "4.070  (SS_res)", "42.055  (SS_tot)"],
            ["R²", "1 − 4.070 / 42.055", "= 0.903"],
        ],
        col_widths=[2.0, 5.5, 9.0]
    )
    note_box(doc,
             "R² = 1.0 → perfect fit.  R² = 0.0 → model is no better than predicting the mean.  "
             "R² < 0.0 → model is worse than predicting the mean (possible for precipitation).")

    # ── 5.5 Paired t-test ─────────────────────────────────────────────────────
    h2(doc, "5.5  Paired t-Test for Systematic Bias")
    body(doc,
         "A two-sided paired t-test is applied to detect systematic over- or under-prediction. "
         "H₀: mean(ŷ − y) = 0. A p-value > 0.05 means predictions are statistically "
         "indistinguishable from actuals — no significant bias.")

    formula(doc,
            "dᵢ  =  ŷᵢ − yᵢ    (prediction error per sample)\n\n"
            "t  =  d̅  /  (sᴰ / √N)    where  sᴰ = std(d),  N = number of samples\n\n"
            "p-value  =  2 × P( tᵏ₋₁ > |t| )")

    h3(doc, "Worked Example (same 5 days)")
    body(doc, "d = [−0.70, −1.40, 1.20, −0.30, −0.30]  →  d̄ = −0.30,  s = 0.868", italic=True)
    body(doc, "t = −0.30 / (0.868 / √5) = −0.30 / 0.388 = −0.773,  df = 4  →  p = 0.48", italic=True)
    body(doc, "Conclusion: p = 0.48 > 0.05 → no statistically significant systematic bias.", italic=True)

    code_block(doc, """\
from scipy import stats

def _paired_ttest(y_true, y_pred):
    t, p = stats.ttest_rel(y_pred, y_true)   # H0: mean(y_pred - y_true) == 0
    return float(t), float(p)""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. VARIABLE IMPORTANCE ANALYSIS
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "6. Variable Importance Analysis  (feature_analysis.py)")
    body(doc,
         "The Variable Importance page answers: "
         "\"Which meteorological variables most significantly influence "
         "the predictive performance of GRU models?\" "
         "The pipeline runs for both stations across all four forecast targets.")

    h2(doc, "6.1  Pipeline Overview")
    data_table(doc,
        ["Step", "Method", "Purpose"],
        [
            ["1", "Coverage filter",          "Remove variables with < threshold% non-NaN data"],
            ["2", "ADF stationarity test",    "Detect unit roots; first-difference non-stationary vars"],
            ["3", "Spearman correlation",     "Measure non-linear rank correlation between all variables"],
            ["4", "Leave-one-out ablation",   "Retrain GRU with one variable removed; measure ΔRMSE"],
            ["5", "Diebold-Mariano test",     "Statistical significance of each variable's contribution"],
            ["6", "Friedman test",            "Non-parametric test across all ablation configurations"],
            ["7", "Benjamini-Hochberg (BH)",  "FDR correction for multiple comparisons"],
            ["8", "Auto conclusion",          "Narrative summary with recommendations"],
        ],
        col_widths=[1.0, 4.5, 11.0]
    )

    # ── 6.2 Coverage ──────────────────────────────────────────────────────────
    h2(doc, "6.2  Coverage Filter")
    formula(doc, "coverageᵥ  =  non_NaN_countᵥ / total_rows")
    body(doc,
         "Variables with coverage below the user-set threshold (default 70%) are "
         "excluded before any analysis. This prevents sparse variables from distorting "
         "the ablation study.")

    h3(doc, "Example — Baguio City Coverage")
    data_table(doc,
        ["Variable", "Coverage (%)", "Action"],
        [
            ["temp  — Avg Temperature",  "99.2%", "✓ Kept"],
            ["pres  — Sea-level Pressure","97.8%","✓ Kept"],
            ["prcp  — Precipitation",    "95.1%", "✓ Kept"],
            ["wspd  — Wind Speed",       "92.4%", "✓ Kept"],
            ["rhum  — Relative Humidity","88.3%", "✓ Kept"],
            ["tsun  — Sunshine Duration","38.2%", "✗ Excluded (< 70%)"],
            ["snow  — Snow Depth",       "12.1%", "✗ Excluded (< 70%)"],
        ],
        col_widths=[5.5, 3.5, 7.5]
    )

    # ── 6.3 ADF Test ──────────────────────────────────────────────────────────
    h2(doc, "6.3  Augmented Dickey-Fuller (ADF) Stationarity Test")
    body(doc,
         "The ADF test checks for a unit root (non-stationarity). A stationary time "
         "series has constant mean and variance — required for valid ARIMA and "
         "correlation analysis.")

    formula(doc,
            "H₀:  unit root present (non-stationary)\n"
            "H₁:  no unit root (stationary)\n\n"
            "Reject H₀ at p < 0.05 → stationary → use as-is\n"
            "Fail to reject H₀ at p ≥ 0.05 → non-stationary → first-difference: Δxₜ = xₜ − xₜ₋₁")

    h3(doc, "Sample ADF Results")
    data_table(doc,
        ["Variable", "ADF Stat", "p-value", "Stationary?", "Action"],
        [
            ["Precipitation (prcp)",    "-12.31", "0.0001", "✓ Yes", "Use as-is"],
            ["Temperature (temp)",      "-3.82",  "0.0028", "✓ Yes", "Use as-is"],
            ["Wind Speed (wspd)",       "-9.57",  "0.0001", "✓ Yes", "Use as-is"],
            ["Sea-level Pressure (pres)","-1.44", "0.1412", "✗ No",  "First-difference"],
        ],
        col_widths=[4.5, 2.5, 2.5, 3.0, 4.0]
    )

    # ── 6.4 Spearman Correlation ───────────────────────────────────────────────
    h2(doc, "6.4  Spearman Rank Correlation")
    body(doc,
         "Spearman correlation is used instead of Pearson because precipitation is "
         "non-normally distributed and many variable relationships are monotonic but "
         "not necessarily linear.")

    formula(doc,
            "rₛ  =  1  −  (6 Σ dᵢ²) / (N(N² − 1))\n\n"
            "where  dᵢ  =  difference in rank between xᵢ and yᵢ for sample i")

    h3(doc, "Interpreting the Correlation Matrix")
    data_table(doc,
        ["ρ Range", "Interpretation"],
        [
            ["|ρ| ≥ 0.8", "Very strong relationship — possible redundancy between variables"],
            ["0.6 ≤ |ρ| < 0.8", "Strong relationship"],
            ["0.4 ≤ |ρ| < 0.6", "Moderate relationship"],
            ["|ρ| < 0.4", "Weak relationship — likely independent contribution"],
        ],
        col_widths=[3.5, 13.0]
    )

    # ── 6.5 Ablation Study ────────────────────────────────────────────────────
    h2(doc, "6.5  Leave-One-Out Ablation Study")
    body(doc,
         "A GRU model is trained with all available variables (baseline). Then, "
         "for each variable, the GRU is retrained with that variable removed. "
         "The change in RMSE measures how much the model degrades without it.")

    formula(doc,
            "ΔRMSEᵥ  =  RMSEₐᵇₗₐₜₑᴰⁿₑᴰ  −  RMSEᵇₐₛₑₗᵢⁿₑ\n\n"
            "ΔRMSE > 0 → removing variable v hurt accuracy → v is important\n"
            "ΔRMSE ≈ 0 → removing v had little effect → v contributes little\n"
            "ΔRMSE < 0 → removing v improved accuracy → v may add noise")

    h3(doc, "Sample Ablation Results — Precipitation @ Baguio City")
    data_table(doc,
        ["Removed Variable", "Baseline RMSE", "Ablated RMSE", "ΔRMSE", "Importance"],
        [
            ["Sea-level Pressure",  "2.41", "3.18", "+0.77", "High"],
            ["Temperature",         "2.41", "2.89", "+0.48", "Moderate"],
            ["Wind Speed",          "2.41", "2.55", "+0.14", "Low"],
            ["Relative Humidity",   "2.41", "2.43", "+0.02", "Negligible"],
            ["Dew Point",           "2.41", "2.38", "-0.03", "Possibly noisy"],
        ],
        col_widths=[4.5, 3.0, 3.0, 2.5, 3.5]
    )

    code_block(doc, """\
def run_ablation_study(df, input_vars, target_col, seq_len, epochs, patience):
    result = {}
    # Baseline: all variables
    result['baseline'] = run_one_ablation(df, input_vars, target_col, ...)

    # Leave-one-out
    result['ablations'] = {}
    for var in input_vars:
        remaining = [v for v in input_vars if v != var]
        result['ablations'][var] = run_one_ablation(df, remaining, target_col, ...)
    return result""")

    # ── 6.6 Diebold-Mariano Test ──────────────────────────────────────────────
    h2(doc, "6.6  Diebold-Mariano Test (HLN-corrected)")
    body(doc,
         "The DM test formally tests whether the ablated model is statistically "
         "significantly worse than the baseline. H₀: removing the variable does "
         "not change forecast accuracy.")

    formula(doc,
            "Loss differential:  dₜ  =  e²₁ₜ  −  e²₂ₜ\n"
            "  where  e₁ = ablated errors,  e₂ = baseline errors\n\n"
            "DM statistic:  DM  =  d̅  /  √(σ̂²ᴰ / T)\n"
            "  (Newey-West long-run variance for σ̂²ᴰ)\n\n"
            "HLN correction factor applied for finite samples\n"
            "Positive DM stat → ablated model (e₁) has larger errors → variable is important")

    # ── 6.7 Friedman Test ─────────────────────────────────────────────────────
    h2(doc, "6.7  Friedman Test")
    body(doc,
         "The Friedman test is a non-parametric alternative to repeated-measures ANOVA. "
         "It tests whether the distributions of block-RMSE differ across all "
         "ablation configurations simultaneously.")
    body(doc,
         "Test errors are split into n_blocks=10 temporal segments. For each block, "
         "RMSE is computed for each model configuration. The Friedman statistic is "
         "computed on the resulting (n_blocks × n_configs) matrix.")
    formula(doc,
            "H₀: all model configurations produce the same median block-RMSE\n"
            "H₁: at least one configuration differs\n\n"
            "Significant (p < 0.05) → at least one variable substantially impacts accuracy")

    # ── 6.8 BH Correction ─────────────────────────────────────────────────────
    h2(doc, "6.8  Benjamini-Hochberg FDR Correction")
    body(doc,
         "When testing K variables simultaneously, the probability of at least one "
         "false positive increases. BH correction controls the False Discovery Rate "
         "at α = 0.05.")

    formula(doc,
            "Procedure:\n"
            "  1. Sort p-values: p₍₁₎ ≤ p₍₂₎ ≤ ... ≤ p₍ᵏ₎\n"
            "  2. Find largest k such that p₍k₎ ≤ (k / K) × α\n"
            "  3. Reject H₀ for all hypotheses 1 ... k")

    h3(doc, "Worked Example — 5 variables tested (α = 0.05)")
    data_table(doc,
        ["Rank (i)", "Variable", "Raw p-value", "BH threshold  (i/K × 0.05)", "Rejected?"],
        [
            ["1", "Sea-level Pressure", "0.003", "0.010", "Yes ✓"],
            ["2", "Temperature",        "0.012", "0.020", "Yes ✓"],
            ["3", "Relative Humidity",  "0.041", "0.030", "No  ✗  (exceeds threshold)"],
            ["4", "Wind Speed",         "0.088", "0.040", "No  ✗"],
            ["5", "Dew Point",          "0.210", "0.050", "No  ✗"],
        ],
        col_widths=[2.0, 4.0, 3.0, 4.5, 3.0]
    )
    note_box(doc,
             "Once a p-value exceeds its BH threshold, all subsequent variables are also "
             "not rejected — even if their p-value is below 0.05. This is stricter than "
             "a naive 0.05 cutoff and reduces false discoveries.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. 7-DAY FORECAST ALGORITHM
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "7. 7-Day Rolling Forecast Algorithm")
    body(doc,
         "After an analysis run, the system can generate a 7-day day-by-day forecast "
         "for any combination of station, model, and variable. The forecast starts "
         "from the day after the last available test record.")

    h2(doc, "7.1  Deep Learning Models (GRU, LSTM, SimpleRNN)")
    data_table(doc,
        ["Step", "Description"],
        [
            ["1", "Load the saved model weights and scaler from cache"],
            ["2", "Initialise last_window = last 45 rows of the test set (scaled), shape (45, 4)"],
            ["3", "For each forecast day d = 1 … 7:"],
            ["3a","  x = last_window[-45:].reshape(1, 45, 4)"],
            ["3b","  pred_scaled = model(x)     →  shape (1, 4)"],
            ["3c","  pred_orig = inverse_transform(pred_scaled)  →  physical units"],
            ["3d","  Append pred_scaled to last_window (window slides forward)"],
            ["4", "Return 7 × 4 forecast array with date labels"],
        ],
        col_widths=[1.5, 15.0]
    )

    code_block(doc, """\
window = test_scaled[-seq_len:].copy()   # (45, 4) seed window

for d in range(7):
    x     = window[-seq_len:].reshape(1, seq_len, n_features)
    pred  = predict_deep(model, x)        # (1, 4) scaled
    forecasts[d] = inverse_transform(pred)
    window = np.vstack([window, pred])    # slide window forward""")

    h2(doc, "7.2  ARIMA")
    body(doc,
         "ARIMA uses the pmdarima .update() and .predict() methods. "
         "The model is updated with the full test set (providing recent context) "
         "then predicts 7 steps ahead in a single call.")

    code_block(doc, """\
for col in columns:
    m = arima_models[col]          # original (not deep-copied here)
    m.update(test_vals[:, j])      # update state with all test observations
    fc = m.predict(n_periods=7)    # 7 bulk predictions (short horizon — stable)
    forecasts[col] = inverse_transform(fc)""")

    h2(doc, "7.3  Linear Regression")
    body(doc,
         "Linear Regression uses the same rolling-window strategy as the deep learning "
         "models. The 45-day window is flattened to (1, 180) for each step.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 8. WEB APPLICATION
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "8. Web Application  (app.py,  pages/)")

    h2(doc, "8.1  Navigation Structure")
    body(doc,
         "The app uses Streamlit's multi-page navigation (st.navigation()). "
         "Two pages are available:")
    bullet(doc, "Model Evaluation  (pages/model_evaluation.py)  — main benchmarking page")
    bullet(doc, "Variable Importance  (pages/2_Variable_Importance.py)  — ablation analysis page")

    h2(doc, "8.2  Model Evaluation Page — Tab Overview")
    data_table(doc,
        ["Tab", "Contents"],
        [
            ["Predictions", "Interactive Plotly 2×2 line chart: Actual vs Predicted for each variable. Hover shows all model values at once. Click legend to toggle models. Drag to zoom, double-click to reset."],
            ["Metrics",     "Numeric table (Accuracy, MSE, MAE, R², t-stat, p-value, Train Time) per variable. Plotly bar charts with 'Zoom to differences' toggle — y-axis starts near min value to amplify small model differences."],
            ["Forecast",    "7-day rolling forecast. Station selector, model checkboxes, variable checkboxes. Output: day-labeled table + multi-model line chart."],
            ["Download",    "Export full predictions and metrics as Excel workbook (.xlsx), one sheet per station."],
        ],
        col_widths=[2.5, 14.0]
    )

    h2(doc, "8.3  Caching System")
    body(doc,
         "Analysis results and trained model weights are persisted to the cache/ directory "
         "so that revisiting the app with the same settings skips retraining entirely.")

    data_table(doc,
        ["File", "Contents", "When Created"],
        [
            ["cache/{key}.pkl",        "station_results dict (actuals, predictions, metrics, train_times, forecast_data)", "After each analysis run"],
            ["cache/{key}_models.pkl", "DL model state_dicts — required for the Forecast tab",                             "After each analysis run"],
            ["cache/{key}_meta.json",  "Human-readable metadata: saved_at, stations, models, variables, date_range",       "After each analysis run"],
        ],
        col_widths=[3.5, 9.5, 3.5]
    )

    h3(doc, "Cache Key Generation")
    body(doc,
         "The key is a 12-character SHA-1 hex digest of all analysis settings. "
         "Any change in station, model list, variables, date range, or hyperparameters "
         "produces a different key and triggers a fresh analysis.")

    code_block(doc, """\
import hashlib

def _make_cache_key(start, end, stations, models, variables,
                    epochs, hidden_size, seq_len, batch_size, patience,
                    num_layers, dropout) -> str:
    s = "|".join([str(start), str(end), ",".join(stations),
                  ",".join(models), ",".join(variables),
                  str(epochs), str(hidden_size), str(seq_len),
                  str(batch_size), str(patience),
                  str(num_layers), str(dropout)])
    return hashlib.sha1(s.encode()).hexdigest()[:12]""")

    h2(doc, "8.4  Variable Importance Page — Tab Overview")
    data_table(doc,
        ["Tab", "Contents"],
        [
            ["① Stationarity", "Variable coverage table + ADF test results per station side-by-side"],
            ["② Correlation",  "Spearman correlation heatmap + ranked correlation tables per target"],
            ["③ Ablation Study","ΔRMSE heatmap (all targets × all variables), baseline GRU metrics, per-target bar chart"],
            ["④ Statistical Tests","Friedman test results per target, DM + BH table per target, raw vs BH p-value comparison chart"],
            ["⑤ Conclusion",   "Cross-station overall variable importance chart (avg ΔRMSE), auto-generated narrative conclusion, per-station breakdowns, CSV download"],
        ],
        col_widths=[3.0, 13.5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 9. HYPERPARAMETER REFERENCE
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "9. Hyperparameter Reference")
    data_table(doc,
        ["Parameter", "Default Value", "Description", "Adjustable in UI?"],
        [
            ["seq_len",     "45",      "Sliding-window lookback in days",                 "Yes"],
            ["hidden_size", "256",     "RNN hidden units per direction",                  "Yes"],
            ["num_layers",  "2",       "Number of stacked RNN layers",                    "Yes"],
            ["dropout",     "0.3",     "Dropout probability between layers",              "Yes"],
            ["batch_size",  "16",      "Mini-batch size for training",                    "Yes"],
            ["epochs",      "200",     "Maximum training epochs",                         "Yes"],
            ["patience",    "30",      "Early stopping patience (val MSE)",               "Yes"],
            ["lr",          "1e-3",    "Adam initial learning rate",                      "No (fixed)"],
            ["LR factor",   "0.5",     "ReduceLROnPlateau multiplier",                    "No (fixed)"],
            ["LR patience", "5",       "Epochs before LR decay",                          "No (fixed)"],
            ["grad clip",   "1.0",     "Maximum gradient L2 norm",                        "No (fixed)"],
            ["MAPE min",    "1.0",     "Minimum |y_true| included in MAPE",               "No (fixed)"],
            ["train split", "80%",     "Temporal train / test ratio",                     "No (fixed)"],
            ["test ratio",  "0.2",     "Fraction of data held out as test set",           "No (fixed)"],
            ["max ARIMA p", "5",       "Maximum AR order in auto_arima search",           "No (fixed)"],
            ["max ARIMA q", "5",       "Maximum MA order in auto_arima search",           "No (fixed)"],
            ["BH α",        "0.05",    "False Discovery Rate threshold (Variable Importance)", "No (fixed)"],
        ],
        col_widths=[3.0, 3.0, 9.5, 2.5]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 10. GPU DEVICE DETECTION
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "10. GPU / Device Detection  (models.py)")
    body(doc,
         "The system auto-detects the best available compute device at import time. "
         "All tensors and models are moved to this device before training.")

    data_table(doc,
        ["Priority", "Device", "Condition", "Typical Speed"],
        [
            ["1st", "CUDA (GPU)",   "torch.cuda.is_available()",              "Fastest — NVIDIA GPU"],
            ["2nd", "MPS  (GPU)",   "torch.backends.mps.is_available()",      "Fast — Apple Silicon"],
            ["3rd", "CPU",          "Fallback",                                "Slowest — any machine"],
        ],
        col_widths=[2.0, 3.5, 6.0, 5.0]
    )

    code_block(doc, """\
def _detect_device():
    if torch.cuda.is_available():
        return torch.device('cuda')
    if hasattr(torch.backends, 'mps') and torch.backends.mps.is_available():
        return torch.device('mps')
    return torch.device('cpu')

DEVICE = _detect_device()   # module-level constant""")

    # ══════════════════════════════════════════════════════════════════════════
    # FOOTER
    # ══════════════════════════════════════════════════════════════════════════
    section_divider(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"University of the Cordilleras  ·  College of CITCS  ·  B.S. Computer Science  ·  {TODAY}"
    )
    r.font.size = Pt(9); r.font.color.rgb = _hex(C_STEEL); r.italic = True

    doc.save("system_documentation.docx")
    print("Saved: system_documentation.docx")


if __name__ == "__main__":
    build()
