"""
patch_conclusions.py
Replaces only paragraphs 56–end (4.1.8 Conclusions + 4.1.9 Recommendations)
in Findings_Objective1_VariableImportance.docx with cross-station-based text.
Everything before paragraph index 56 is left completely untouched.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Load document ────────────────────────────────────────────────────────────
doc = Document('Findings_Objective1_VariableImportance.docx')

# ── Locate cut point (first paragraph of 4.1.8 Conclusions) ─────────────────
CUT_INDEX = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading') and '4.1.8' in p.text:
        CUT_INDEX = i
        break

if CUT_INDEX is None:
    raise RuntimeError("Could not find '4.1.8 Conclusions' heading.")
print(f"Cut point found at paragraph index {CUT_INDEX}")

# ── Remove all paragraphs from CUT_INDEX onward ──────────────────────────────
# python-docx paragraphs share the same parent body element
body = doc.element.body
all_paras = doc.paragraphs[:]  # snapshot

# Collect the XML elements to remove
elements_to_remove = []
for p in all_paras[CUT_INDEX:]:
    elements_to_remove.append(p._element)

for el in elements_to_remove:
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)

print(f"Removed {len(elements_to_remove)} elements from index {CUT_INDEX} onward.")


# ── Helper: add body paragraph ───────────────────────────────────────────────
def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    return p


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


# ═══════════════════════════════════════════════════════════════════════════════
#  4.1.8 CONCLUSIONS  (cross-station data-driven)
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4.1.8 Conclusions', level=2)

add_body(doc,
    'The following conclusions are drawn from the cross-station combined analysis of variable '
    'importance, integrating results from the leave-one-out ablation study, Diebold-Mariano '
    '(DM) significance tests with Harvey-Leybourne-Newbold correction, Friedman non-parametric '
    'ranking tests, and Benjamini-Hochberg FDR correction applied across six '
    'target/station combinations—three forecast targets (temperature, wind speed, and pressure) '
    'evaluated at two stations (Baguio and Manila). The analysis addresses the first research '
    'objective: Which meteorological variables most significantly influence the predictive '
    'performance of GRU models in weather forecasting?'
)

add_body(doc,
    'Conclusion 1: Sea-level pressure is the most consistently influential meteorological '
    'variable across both stations and all forecast targets. Among the four retained input '
    'variables, sea-level pressure produced the highest average ΔRMSE when removed from the '
    'GRU input set (average ΔRMSE = +0.3239 across all six target/station combinations), and '
    'was statistically significant in two of the six combinations after Benjamini-Hochberg '
    'correction. This result reflects the role of pressure as a fundamental organizing variable '
    'in atmospheric dynamics: pressure gradient forces drive wind, influence moisture transport, '
    'and are causally linked to temperature advection patterns—particularly during synoptic-scale '
    'weather events such as typhoons, cold surges, and monsoon transitions that characterize '
    'the Philippine climate calendar. The elevated average ΔRMSE of pressure confirms that its '
    'removal degrades GRU accuracy more than the removal of any other cross-predictor, even in '
    'configurations where the degradation does not individually meet the formal significance '
    'threshold after multiple-comparison correction.'
)

add_body(doc,
    'Conclusion 2: Wind speed is the most consistently significant cross-predictor, reaching '
    'statistical significance in four of the six target/station combinations after '
    'Benjamini-Hochberg correction—more than any other variable in the study. Wind speed was '
    'significant at Baguio for pressure prediction (DM stat = 3.14, BH p = 0.0036), reflecting '
    'the geostrophic relationship between surface wind and pressure gradient in highland terrain. '
    'At Manila, wind speed reached significance for temperature prediction (DM stat = 3.14, '
    'BH p = 0.0036), consistent with the coastal sea-breeze circulation that couples wind speed '
    'and temperature diurnally and seasonally. The average ΔRMSE for wind speed across all '
    'combinations (+0.1275) ranks third overall, but its frequency of significance across '
    'stations and targets makes it the most reliably useful cross-predictor in the GRU input set. '
    'These findings establish wind speed as a high-priority variable for inclusion in '
    'multi-parameter GRU weather forecasting models in the Philippine context.'
)

add_body(doc,
    'Conclusion 3: The cross-station significance pattern differs meaningfully between Baguio '
    'and Manila, confirming that variable importance in GRU forecasting is inherently '
    'location-specific. At Baguio, three variables were identified as significant across the '
    'three forecast targets: average temperature (its own lag for temperature prediction), '
    'sea-level pressure (its own lag for pressure prediction), and wind speed (as a predictor '
    'of pressure). At Manila, four variables were significant: average temperature, sea-level '
    'pressure, wind speed, and—uniquely to Manila—relative humidity (significant for temperature '
    'prediction, DM stat = 3.02, BH p = 0.0036). The appearance of relative humidity as a '
    'significant predictor only at Manila—and not at Baguio—reflects the stronger humidity-'
    'temperature coupling at coastal sea-level locations, where evaporative feedback and maritime '
    'air mass characteristics create tighter thermodynamic linkages between these variables than '
    'are present in the orographically disrupted highland environment of Baguio. This '
    'location-dependence precludes the application of a universal variable importance ranking '
    'across Philippine weather stations and necessitates site-specific ablation analysis.'
)

add_body(doc,
    'Conclusion 4: Relative humidity exhibits the weakest cross-predictor influence '
    'when evaluated across all target/station combinations, with an average ΔRMSE of −0.0081 '
    '(indicating that its removal marginally improves the average baseline performance '
    'across most configurations) and statistical significance in only one of six combinations. '
    'This finding should not be interpreted as evidence that relative humidity is unimportant '
    'in weather forecasting broadly—rather, it reflects that in the specific GRU architecture '
    'and input representation used in this study, humidity\'s predictive signal for temperature, '
    'wind speed, and pressure is largely subsumed by the autoregressive information already '
    'present in those variables\' own lags, with the exception of Manila\'s temperature '
    'prediction. The slight negative average ΔRMSE suggests that, in some configurations, '
    'humidity introduces marginal multicollinearity noise that slightly degrades model '
    'performance—a known risk when adding highly correlated inputs to recurrent networks '
    'with finite training sequences (Goodfellow et al., 2016).'
)

add_body(doc,
    'Conclusion 5: The Friedman test confirmed that the collective effect of variable subset '
    'configuration on GRU forecast performance is statistically significant across all six '
    'target/station combinations (Friedman statistics ranging from 14.24 to 27.20, all '
    'p < 0.01). This global test result validates the premise of the ablation pipeline: '
    'input feature selection meaningfully affects GRU weather forecast quality, and the '
    'observed performance differences are not attributable to random sampling variability. '
    'The combined evidence from DM pairwise tests and the Friedman global test provides '
    'a multi-level statistical foundation for the variable importance conclusions presented '
    'above, addressing both the inferential question of whether differences are real and the '
    'practical question of which specific variables drive those differences.'
)


# ═══════════════════════════════════════════════════════════════════════════════
#  4.1.9 RECOMMENDATIONS  (cross-station data-driven)
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4.1.9 Recommendations', level=2)

add_body(doc,
    'Based on the cross-station combined conclusions of the variable importance analysis, the '
    'following recommendations are directed at researchers and practitioners developing GRU-based '
    'daily weather forecasting systems for Philippine climate stations. The recommendations '
    'are ordered by priority, reflecting the magnitude and consistency of each variable\'s '
    'influence as quantified by average ΔRMSE and frequency of statistical significance '
    'across the six target/station combinations evaluated.'
)

add_body(doc,
    'Recommendation 1: Prioritize sea-level pressure as a mandatory input in GRU weather '
    'forecasting models. With the highest average ΔRMSE (+0.3239) across all combinations, '
    'sea-level pressure is the single variable whose removal produces the largest average '
    'forecast degradation. This priority holds regardless of which meteorological variable '
    'is being predicted—pressure contains synoptic-scale information about weather system '
    'organization that benefits the prediction of co-varying atmospheric quantities. In '
    'operational deployment, pressure data should be verified for continuity and sensor '
    'calibration, as even small systematic biases in pressure readings can propagate into '
    'GRU predictions given the variable\'s dominant role in the input representation.'
)

add_body(doc,
    'Recommendation 2: Include wind speed as a cross-predictor in all multi-variable GRU '
    'configurations. Wind speed achieved the highest frequency of statistical significance '
    'across target/station combinations (4 of 6), and its average ΔRMSE of +0.1275 ranks '
    'third among all variables. Its significance for both pressure prediction at Baguio and '
    'temperature prediction at Manila demonstrates that wind speed carries meteorologically '
    'meaningful cross-variable information that GRU models can exploit regardless of which '
    'target variable is being forecasted. Researchers should ensure that wind speed measurements '
    'meet WMO standards for anemometer height (10 m) and that station exposure conditions '
    'are documented, as orographic distortions in highland stations can reduce the '
    'representativeness of wind speed measurements.'
)

add_body(doc,
    'Recommendation 3: Conduct station-specific assessments to determine whether '
    'relative humidity should be included as a GRU input. The cross-station analysis '
    'revealed that relative humidity reaches significance only at Manila (for temperature '
    'prediction) and not at Baguio. Before including humidity as an input at additional '
    'stations, site-specific ablation analysis should confirm whether a statistically '
    'significant contribution can be established at that location. At stations with strong '
    'maritime influence or urban heat island effects, the humidity-temperature coupling that '
    'drives Manila\'s significance result is more likely to replicate. At complex terrain '
    'stations, the multicollinearity between humidity and pressure may instead cause '
    'humidity to degrade rather than improve GRU performance, consistent with its '
    'average ΔRMSE of −0.0081 across all stations.'
)

add_body(doc,
    'Recommendation 4: Apply the cross-station ablation-DM-Friedman pipeline as a '
    'standardized variable selection protocol before deploying GRU weather models at new '
    'Philippine stations. The three-stage pipeline demonstrated in this study—ablation for '
    'practical effect size estimation, DM testing for pairwise statistical significance, '
    'and Friedman testing for global configuration differentiation—provides a rigorous and '
    'reproducible framework for input variable selection. This protocol should be applied '
    'at additional PAGASA monitoring stations across the Philippine archipelago, especially '
    'in regions with distinct climatological regimes: the Eastern Visayas (Leyte, Samar), '
    'where typhoon landfall frequency is highest; the Cagayan Valley, which experiences '
    'distinct Amihan and Habagat dynamics; and the island of Mindanao, which is '
    'climatologically distinct from Luzon and exhibits different monsoon seasonality. '
    'Standardizing this protocol across the network would enable a national-level '
    'understanding of geographic variability in GRU feature importance.'
)

add_body(doc,
    'Recommendation 5: Investigate the temporal variability of variable importance by '
    'conducting seasonally stratified ablation studies. The current analysis treats '
    'variable importance as a stationary property over the 2020–2026 evaluation period. '
    'However, the physical mechanisms linking meteorological variables—pressure-wind '
    'geostrophy, humidity-temperature coupling, wind-pressure cross-prediction—are inherently '
    'season-dependent. During the southwest monsoon season (June–September), when moisture '
    'flux from the South China Sea dominates, relative humidity-temperature coupling is '
    'likely stronger. During the northeast monsoon (November–February), pressure-wind '
    'coupling may be more pronounced due to enhanced large-scale pressure gradients. '
    'Future studies should segment the ablation analysis by season and by ENSO phase '
    '(El Niño vs. La Niña) to characterize how GRU variable importance shifts with '
    'synoptic regime—findings that would directly support the development of '
    'season-adaptive or regime-switching GRU architectures for operational Philippine '
    'weather forecasting.'
)

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save('Findings_Objective1_VariableImportance_UPDATED.docx')
print("Saved: Findings_Objective1_VariableImportance_UPDATED.docx — conclusions and recommendations updated.")
